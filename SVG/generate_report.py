#!/usr/bin/env python3
"""
Generate comprehensive Mental Health Support Platform project report as DOCX.
Embeds all SVG diagrams as PNG images inside the Word document.
"""

import os
import io
import glob
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

BASE_DIR = r"c:\Users\mdsha\Downloads\SVG\SVG"
WORD_DIR = os.path.join(BASE_DIR, "Word")
PNG_CACHE_DIR = os.path.join(WORD_DIR, "_png_cache")
os.makedirs(WORD_DIR, exist_ok=True)
os.makedirs(PNG_CACHE_DIR, exist_ok=True)

# ─── Styling helpers ───

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="1a5276"):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EBF5FB")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def convert_all_svgs_to_png():
    """Pre-convert all SVGs to high-res PNG using Playwright headless browser."""
    from playwright.sync_api import sync_playwright
    svg_files = glob.glob(os.path.join(BASE_DIR, "*.svg"))
    print(f"  Converting {len(svg_files)} SVG files to high-res PNG (3x scale)...")

    # Clear old cached PNGs to force re-render at higher quality
    for old_png in glob.glob(os.path.join(PNG_CACHE_DIR, "*.png")):
        os.remove(old_png)

    SCALE = 3  # 3x device pixel ratio for crisp images

    with sync_playwright() as p:
        browser = p.chromium.launch()
        context = browser.new_context(device_scale_factor=SCALE)
        page = context.new_page()

        for svg_path in svg_files:
            png_name = os.path.basename(svg_path).replace('.svg', '.png')
            png_path = os.path.join(PNG_CACHE_DIR, png_name)
            try:
                svg_content = open(svg_path, 'r', encoding='utf-8').read()
                page.set_content(f"""<!DOCTYPE html>
<html><head><style>
body {{ margin: 0; padding: 0; background: white; display: inline-block; }}
svg {{ max-width: none; }}
</style></head>
<body>{svg_content}</body></html>""")
                page.wait_for_timeout(500)
                # Get the SVG element's bounding box
                svg_el = page.query_selector('svg')
                if svg_el:
                    box = svg_el.bounding_box()
                    if box and box['width'] > 0 and box['height'] > 0:
                        vw = max(int(box['width']) + 40, 1200)
                        vh = max(int(box['height']) + 40, 800)
                        page.set_viewport_size({'width': vw, 'height': vh})
                        page.wait_for_timeout(300)
                        svg_el.screenshot(path=png_path)
                        print(f"    ✓ {os.path.basename(svg_path)}")
                        continue
                # Fallback: full page screenshot
                page.screenshot(path=png_path, full_page=True)
                print(f"    ✓ {os.path.basename(svg_path)} (full page)")
            except Exception as e:
                print(f"    ✗ {os.path.basename(svg_path)}: {e}")

        context.close()
        browser.close()
    print("  SVG conversion complete (high-res).")


def svg_to_png_bytes(svg_path, width=None):
    """Get PNG bytes for a pre-converted SVG."""
    png_name = os.path.basename(svg_path).replace('.svg', '.png')
    png_path = os.path.join(PNG_CACHE_DIR, png_name)
    if os.path.exists(png_path):
        with open(png_path, 'rb') as f:
            return f.read()
    return None


def add_svg_image(doc, svg_path, caption=None, width_inches=6.0):
    """Add SVG as PNG image to doc. Falls back to placeholder if conversion fails."""
    png_data = svg_to_png_bytes(svg_path, width=int(width_inches * 150))
    if png_data:
        stream = io.BytesIO(png_data)
        doc.add_picture(stream, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Fallback: add a note
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Diagram: {os.path.basename(svg_path).replace('.svg','')}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.size = Pt(10)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("(Please refer to the attached SVG file for the full diagram)")
        run2.italic = True
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        cap.space_after = Pt(6)


def add_heading_styled(doc, text, level=1):
    """Add heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    return heading


def add_body_text(doc, text, bold=False, italic=False, size=11):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    return p


def add_bullet(doc, text, level=0):
    """Add bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ─── MAIN DOCUMENT GENERATION ───

def generate_report():
    print("Creating Mental Health Support Platform Report...")
    print("Step 1: Converting SVG diagrams to PNG images...")
    convert_all_svgs_to_png()
    print("Step 2: Building DOCX document...")
    doc = Document()

    # ─── Page Setup ───
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ─── Default font ───
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # ═══════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MENTAL HEALTH SUPPORT PLATFORM")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Software Requirements & System Design Report")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    doc.add_paragraph()

    # Horizontal line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("━" * 60)
    run_line.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
    run_line.font.size = Pt(10)

    doc.add_paragraph()

    # Project info
    info_items = [
        ("Project Type:", "Web Application (Full Stack)"),
        ("Duration:", "12 Weeks"),
        ("Institution:", "Kent Institute of Technology Australia"),
        ("Industry Sponsor:", "Skillup Labs"),
        ("Team:", "Mel23 Tech Solution"),
        ("Date:", "April 2026"),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(label + " ")
        run_l.bold = True
        run_l.font.size = Pt(12)
        run_l.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        run_v = p.add_run(value)
        run_v.font.size = Pt(12)
        run_v.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(3):
        doc.add_paragraph()

    # Team members on cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Team Members")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    members = [
        "Sanjana Tanwar (K240225) — Team Leader / Project Manager",
        "Jubair Zaman Dipto (K240568) — Full-Stack / Tech Lead",
        "S M Nalid Maola (K231528) — Backend Developer",
        "Nabil Ashrafi (K231720) — Frontend Developer",
        "Md Al Amin Sikder (K240094) — QA Engineer / Tester",
        "Tranpreet Kaur Kumar (K241957) — UI/UX Designer",
    ]
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(m)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Academic Supervisor: Mr. Syed Altaf")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Industry Sponsor Contact: Nabin Singh (wil@skilluplabs.com.au)")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # TABLE OF CONTENTS (placeholder)
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_entries = [
        ("1.", "Introduction", ""),
        ("  1.1", "Purpose of the Project", ""),
        ("  1.2", "Objective of the Project", ""),
        ("  1.3", "Project Scope and Limitations", ""),
        ("  1.4", "Assumptions", ""),
        ("  1.5", "Company Profile", ""),
        ("  1.6", "Organisation Chart", ""),
        ("  1.7", "Roles and Responsibilities", ""),
        ("  1.8", "Key Deliverables", ""),
        ("  1.9", "Project Budget", ""),
        ("  1.10", "Work Breakdown Structure", ""),
        ("  1.11", "Gantt Chart", ""),
        ("2.", "System Requirements", ""),
        ("  2.1", "Functional Requirements", ""),
        ("  2.2", "Non-Functional Requirements", ""),
        ("3.", "System Architecture", ""),
        ("  3.1", "Software and Hardware Architecture", ""),
        ("  3.2", "Tech Stack (Detailed)", ""),
        ("4.", "System Design", ""),
        ("  4.1", "User Interface Design / Wireframes", ""),
        ("5.", "Use Case Diagrams", ""),
        ("6.", "Context Diagrams (Gane & Sarson Method)", ""),
        ("7.", "Data Flow Diagrams (Gane & Sarson Method)", ""),
        ("  7.1", "DFD Level 1", ""),
        ("  7.2", "DFD Level 2", ""),
        ("8.", "Sequence Diagrams", ""),
        ("9.", "Entity Relationship Diagram & Data Dictionary", ""),
        ("10.", "Class Diagram", ""),
        ("11.", "Data Protection Impact Assessment (DPIA)", ""),
        ("12.", "Test and Implementation Plan", ""),
        ("13.", "Responses on Feedback from Supervisor", ""),
        ("", "References", ""),
        ("", "Appendix", ""),
    ]
    for num, title, _ in toc_entries:
        p = doc.add_paragraph()
        if num.startswith("  "):
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(f"{num.strip()} {title}")
            run.font.size = Pt(10)
        else:
            run = p.add_run(f"{num} {title}")
            run.bold = True
            run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "1. Introduction", level=1)

    add_body_text(doc, "Mental health is a critical component of overall wellbeing, yet many individuals face challenges in accessing timely support, tracking their emotional state, and finding reliable mental health resources. Barriers such as stigma, lack of awareness, and limited access to support systems often prevent individuals from seeking help.")
    add_body_text(doc, "This document presents the comprehensive Software Requirements Specification (SRS) and System Design report for the Mental Health Support Platform — a web-based application designed to enable users to track their mood, access curated mental health resources, and engage with supportive tools in a secure and user-friendly environment.")

    # 1.1 Purpose
    add_heading_styled(doc, "1.1 Purpose of the Project", level=2)
    add_body_text(doc, "The purpose of this project is to design and create a safe and easy-to-use web-based Mental Health Support Platform, which allows users to monitor and manage their mental health. The system will offer basic features that will include mood tracking, journaling, and access to curated mental health resources in a private and secure setting.")
    add_body_text(doc, "The goal of the platform is to help in self-awareness and early intervention, and it is expected to assist users to document, track, and analyze their emotional patterns over time. The feature will allow users to record their moods on a daily basis, keep personal journal entries, and provide organized materials such as articles, tips, and support information on mental health.")
    add_body_text(doc, "The system will be designed with secure authentication and data protection that will guarantee confidentiality, integrity, and safe management of sensitive user information. Also, the platform will show user information in a straightforward dashboard to enhance usability and interaction.")
    add_body_text(doc, "The project will follow Software Development Lifecycle (SDLC) practices, including requirement analysis, system design, development, testing, and deployment to provide a functional and reliable prototype of the project.")

    # 1.2 Objective
    add_heading_styled(doc, "1.2 Objective of the Project", level=2)
    add_body_text(doc, "The objective of this project is to design and develop a prototype platform that supports users in monitoring their mental wellbeing and accessing helpful resources. The key objectives include:")
    objectives = [
        "Analyse user needs related to mental health tracking and support",
        "Define system requirements for a secure support platform",
        "Design and implement a full-stack web application using React, Node.js and PostgreSQL",
        "Develop mood tracking, journaling, and AI chatbot features",
        "Implement secure authentication (JWT, OTP, bcrypt) and data protection mechanisms (AES-256)",
        "Provide access to curated mental health resources with categorization",
        "Enable real-time secure messaging between patients and clinicians",
        "Integrate appointment scheduling with Google Calendar synchronization",
        "Deploy a working prototype system on a cloud hosting environment",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    # 1.3 Project Scope and Limitations
    add_heading_styled(doc, "1.3 Project Scope and Limitations", level=2)
    add_body_text(doc, "Scope:", bold=True)
    scope_items = [
        "User registration, authentication (JWT + OTP), and role-based access control (RBAC)",
        "Daily mood tracking with 1–10 scale, notes, and trend visualization",
        "Secure messaging with AES-256 encryption between patients and clinicians",
        "Appointment scheduling with slot conflict detection and Google Calendar sync",
        "AI-powered chatbot with NLP processing and clinician escalation (confidence < 70%)",
        "Emergency support button with < 1 second response time and crisis hotline display",
        "Clinical records management with consent-based access control",
        "Resource library with URL validation, categorization, and view tracking",
        "Administrative dashboard with analytics, audit trails, and compliance reporting",
        "Privacy controls including consent management and right to data deletion",
    ]
    for item in scope_items:
        add_bullet(doc, item)

    add_body_text(doc, "Limitations:", bold=True)
    limitations = [
        "The platform is a prototype and not intended for clinical diagnosis or treatment",
        "AI chatbot responses are supplementary and do not replace professional medical advice",
        "Real-time video consultation features are outside the current scope",
        "The system will be tested with simulated data, not real patient records",
        "Mobile-native application is out of scope (responsive web design will be provided)",
    ]
    for item in limitations:
        add_bullet(doc, item)

    # 1.4 Assumptions
    add_heading_styled(doc, "1.4 Assumptions", level=2)
    assumptions = [
        "Users will have access to a modern web browser (Chrome, Firefox, Safari, Edge)",
        "Users have a stable internet connection for real-time features",
        "The development team has access to required development tools and cloud services",
        "Third-party APIs (Google Calendar, OpenAI, Twilio, SendGrid) will remain available during development",
        "The project will follow agile methodology with iterative development sprints",
        "Team members will be available throughout the 12-week project duration",
        "The academic supervisor and industry sponsor will provide timely feedback",
        "PostgreSQL database will be hosted on a cloud provider with automated backups",
    ]
    for a in assumptions:
        add_bullet(doc, a)

    # 1.5 Company Profile
    add_heading_styled(doc, "1.5 Company Profile", level=2)
    add_body_text(doc, "Mel23 Tech Solution", bold=True, size=13)
    add_body_text(doc, "Mel23 Tech Solution is a student-led development team formed at Kent Institute of Technology Australia for the WIL (Work Integrated Learning) capstone project. The team operates under the mentorship of the academic supervisor, Mr. Syed Altaf, and with the industry sponsorship of Skillup Labs.")
    add_body_text(doc, "Industry Sponsor: Skillup Labs", bold=True)
    add_body_text(doc, "Skillup Labs is the industry project provider that sponsors the Mental Health Support Platform project. The organization supports student teams in building real-world technology solutions within the HealthTech domain.")

    info_table = [
        ["Industry Provider", "Skillup Labs"],
        ["Program Contact", "Nabin Singh"],
        ["Contact Email", "wil@skilluplabs.com.au"],
        ["Academic Supervisor", "Mr. Syed Altaf"],
        ["Institution", "Kent Institute of Technology Australia"],
        ["Project Duration", "12 Weeks"],
        ["Team Size", "6 Members"],
    ]
    add_styled_table(doc, ["Attribute", "Details"], info_table, col_widths=[2.5, 4.0])

    # 1.6 Organisation Chart
    add_heading_styled(doc, "1.6 Organisation Chart", level=2)
    add_body_text(doc, "The following organisation chart illustrates the project team hierarchy, including the academic supervisor, industry sponsor, team leader, and individual team members with their designated roles.")
    svg_path = os.path.join(BASE_DIR, "1.6 Organisation chart_diagram_1.svg")
    add_svg_image(doc, svg_path, "Figure 1.1: Organisation Chart — Mel23 Tech Solution", 6.2)

    # 1.7 Roles and Responsibilities
    add_heading_styled(doc, "1.7 Roles and Responsibilities", level=2)
    roles_data = [
        ["Mr. Syed Altaf", "Academic Supervisor", "Provides academic guidance, reviews deliverables, evaluates team progress"],
        ["Nabin Singh", "Industry Sponsor", "Provides industry context, project requirements, and real-world expectations"],
        ["Sanjana Tanwar\nK240225", "Team Leader /\nProject Manager", "Project coordination, sprint planning, stakeholder communication, progress tracking"],
        ["Jubair Zaman Dipto\nK240568", "Full-Stack /\nTech Lead", "System architecture, API design, full-stack integration, code reviews"],
        ["S M Nalid Maola\nK231528", "Backend Developer", "REST API development, database design, business logic, authentication"],
        ["Nabil Ashrafi\nK231720", "Frontend Developer", "React component development, responsive UI, state management"],
        ["Md Al Amin Sikder\nK240094", "QA Engineer / Tester", "Test planning, test automation, bug tracking, CI/CD pipeline"],
        ["Tranpreet Kaur Kumar\nK241957", "UI/UX Designer", "Wireframes, UX research, accessibility compliance, design documentation"],
    ]
    add_styled_table(doc, ["Name / ID", "Role", "Responsibilities"], roles_data, col_widths=[1.8, 1.5, 3.5])

    # 1.8 Key Deliverables
    add_heading_styled(doc, "1.8 Key Deliverables", level=2)
    deliverables = [
        ["D1", "Problem Analysis Report", "Week 2"],
        ["D2", "Software Requirements Specification (SRS)", "Week 3"],
        ["D3", "System Architecture Documentation", "Week 4"],
        ["D4", "Database Schema Design (ERD)", "Week 5"],
        ["D5", "UI/UX Wireframes (Balsamiq Cloud)", "Week 5"],
        ["D6", "Functional Web Application Prototype", "Week 9"],
        ["D7", "Source Code Repository (GitHub)", "Ongoing"],
        ["D8", "Testing Documentation (Unit, Integration, E2E)", "Week 11"],
        ["D9", "Deployment Documentation", "Week 11"],
        ["D10", "Final Presentation & Demonstration", "Week 12"],
    ]
    add_styled_table(doc, ["ID", "Deliverable", "Due"], deliverables, col_widths=[0.6, 3.5, 1.2])

    # 1.9 Project Budget
    add_heading_styled(doc, "1.9 Project Budget", level=2)
    add_body_text(doc, "As this is an academic capstone project, the budget is allocated for cloud services, third-party API usage, and development tools. The following table outlines the estimated project budget:")
    budget = [
        ["Cloud Hosting (AWS/Render)", "$0 – $50/month", "Free tier / student credits"],
        ["PostgreSQL Database Hosting", "$0 – $20/month", "Render/Supabase free tier"],
        ["Domain Name", "$10 – $15/year", "Optional for demo"],
        ["OpenAI API (Chatbot)", "$10 – $30/month", "GPT-3.5/4 usage during development"],
        ["Twilio SMS API", "$0 – $15/month", "Trial credits available"],
        ["SendGrid Email API", "$0", "Free tier (100 emails/day)"],
        ["Google Calendar API", "$0", "Free tier"],
        ["Balsamiq Cloud (Wireframes)", "$0", "Educational license"],
        ["GitHub (Version Control)", "$0", "Free for students"],
        ["Development Tools (VS Code etc.)", "$0", "Free / open-source"],
        ["Total Estimated Budget", "$20 – $130/month", ""],
    ]
    add_styled_table(doc, ["Item", "Estimated Cost", "Notes"], budget, col_widths=[2.5, 1.5, 2.5])

    # 1.10 Work Breakdown Structure
    add_heading_styled(doc, "1.10 Work Breakdown Structure", level=2)
    add_body_text(doc, "The Work Breakdown Structure (WBS) decomposes the project into six major phases: Discovery, Requirements, Design, Development, Testing, and Deployment. Each phase contains specific tasks and sub-deliverables as illustrated below.")
    svg_path = os.path.join(BASE_DIR, "1.10 Work breakdown structure_diagram_1.svg")
    add_svg_image(doc, svg_path, "Figure 1.2: Work Breakdown Structure — Mental Health Support Platform", 6.5)

    # WBS table
    wbs_data = [
        ["1. Discovery\n(Week 1–2)", "Problem Statement, Market Research, User Personas, Project Vision, Competitor Analysis"],
        ["2. Requirements\n(Week 3–4)", "Functional Requirements (75), Non-Functional Requirements, Use Case Diagrams, User Stories, SRS Document, System Architecture"],
        ["3. Design\n(Week 5–6)", "Database Schema, API Specification, ERD Design, Wireframes (UI/UX), DFD (Gane & Sarson), Class Diagram"],
        ["4. Development\n(Week 7–9)", "Backend (Express APIs), Frontend (React UI), Database Implementation, Auth & Security, AI Chatbot (OpenAI), Real-time Messaging, Feature Integration"],
        ["5. Testing\n(Week 10–11)", "Unit Testing (Jest), Integration Testing, E2E Testing (Cypress), Security Testing, User Acceptance Testing"],
        ["6. Deployment\n(Week 11–12)", "Cloud Deployment, CI/CD Pipeline, Final Report, Demo Presentation, Documentation"],
    ]
    add_styled_table(doc, ["Phase / Timeline", "Tasks"], wbs_data, col_widths=[1.8, 5.0])

    # 1.11 Gantt Chart
    add_heading_styled(doc, "1.11 Gantt Chart", level=2)
    add_body_text(doc, "The following Gantt chart provides a timeline view of the project phases and major milestones across the 12-week development period.")

    gantt_data = [
        ["Week 1–2", "Discovery & Planning", "Problem research, personas, vision", "✓"],
        ["Week 3", "Requirements Documentation", "SRS, functional/non-functional requirements", "✓"],
        ["Week 4", "System Architecture Design", "Architecture diagrams, tech stack decisions", "✓"],
        ["Week 5", "Database & API Design", "ERD, API spec, DFD diagrams", "✓"],
        ["Week 6", "UI/UX Design", "Wireframes (Balsamiq), design review", "✓"],
        ["Week 7", "Backend Development Pt. 1", "Auth, user management, appointments APIs", "In Progress"],
        ["Week 8", "Backend Development Pt. 2", "Mood, messaging, chatbot, emergency APIs", "In Progress"],
        ["Week 9", "Frontend Development", "React UI, dashboard, all page components", "Planned"],
        ["Week 10", "System Integration", "Frontend-backend integration, real-time features", "Planned"],
        ["Week 11", "Testing & Deployment", "Unit, integration, E2E, security testing, deploy", "Planned"],
        ["Week 12", "Final Demo & Report", "Presentation, documentation, handover", "Planned"],
    ]
    add_styled_table(doc, ["Period", "Phase", "Key Activities", "Status"], gantt_data, col_widths=[1.0, 2.0, 3.0, 0.8])

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 2. SYSTEM REQUIREMENTS
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "2. System Requirements", level=1)
    add_body_text(doc, "This section defines the functional and non-functional requirements of the Mental Health Support Platform. The requirements have been elicited through analysis of the project proposal, stakeholder consultation, and industry best practices in HealthTech application development.")

    # 2.1 Functional Requirements
    add_heading_styled(doc, "2.1 Functional Requirements", level=2)
    add_body_text(doc, "The platform defines 75 functional requirements categorized across Authentication, Profile Management, Appointments, Clinical Records, Messaging, Mood Tracking, AI Chatbot, Emergency Support, Notifications, Content Management, Analytics, Administration, System Maintenance, Privacy, and UI/Localization modules.")

    # FR table - full list
    fr_data = [
        ["FR1", "Register new user account with unique email/SIM validation", "Create", "Authentication"],
        ["FR2", "Verify accounts via OTP within a 5-second window", "Create", "Authentication"],
        ["FR3", "Lock user account after 5 failed login attempts", "Update", "Authentication"],
        ["FR4", "Record user login timestamp and IP address", "Create", "Audit"],
        ["FR5", "Process password reset with 10-minute token validity", "Update", "Authentication"],
        ["FR6", "Update profile fields with strict input validation", "Update", "Profile"],
        ["FR7", "Upload and store profile image (size < 2MB)", "Create", "Profile"],
        ["FR8", "User Login/Logout", "Read", "Authentication"],
        ["FR9", "Schedule appointment with time-slot conflict detection", "Create", "Appointment"],
        ["FR10", "Reject booking if timeslot is unavailable", "Read", "Appointment"],
        ["FR11", "Reschedule appointments up to 2 hours before session", "Update", "Appointment"],
        ["FR12", "Mark appointment as missed after 15-minute delay", "Update", "Appointment"],
        ["FR13", "Dispatch booking confirmation within 1 minute", "Create", "Notification"],
        ["FR14", "Approve pending appointments within 10-minute window", "Update", "Appointment"],
        ["FR15", "Auto-cancel pending bookings after 30 minutes inactivity", "Delete", "Appointment"],
        ["FR16", "Record clinical session notes post-consultation", "Create", "Clinical"],
        ["FR17", "Restrict session note access based on consent flags", "Read", "Security"],
        ["FR18", "View comprehensive patient history (authorized only)", "Read", "Clinical"],
        ["FR19", "Retain patient records independently of clinician employment", "Read", "Compliance"],
        ["FR20", "Reassign patients upon clinician departure", "Update", "Clinical"],
        ["FR21", "Request access to previous medical records", "Read", "Clinical"],
        ["FR22", "Log all medical record access and viewing actions", "Create", "Audit"],
        ["FR23", "Encrypt all messages using AES-256", "Create", "Security"],
        ["FR24", "Deliver chat messages with < 2 seconds latency", "Create", "Messaging"],
        ["FR25", "Store chat transcripts permanently in database", "Create", "Database"],
        ["FR26", "Delete chat messages locally (soft delete only)", "Delete", "Messaging"],
        ["FR27", "Submit daily mood entry with exact timestamp", "Create", "Mood"],
        ["FR28", "Prevent duplicate mood entries within 1-hour timeframe", "Create", "Mood"],
        ["FR29", "Generate automated weekly mood reports", "Read", "Mood"],
        ["FR30", "Trigger risk alert if mood < 2 for 3 consecutive days", "Create", "AI"],
        ["FR31", "Recommend resources based on mood patterns", "Read", "AI"],
        ["FR32", "Respond to user queries via Chatbot within 3 seconds", "Create", "AI Chatbot"],
        ["FR33", "Escalate Chatbot session if confidence < 70%", "Update", "AI Chatbot"],
        ["FR34", "Record all Chatbot conversation transcripts", "Create", "AI Chatbot"],
        ["FR35", "Review Chatbot logs via admin dashboard", "Read", "Admin"],
        ["FR36", "Identify users with 7 days of inactivity", "Read", "AI"],
        ["FR37", "Dispatch automated wellness check notifications", "Create", "Notification"],
        ["FR38", "Activate emergency button with < 1 second response", "Create", "Emergency"],
        ["FR39", "Display crisis hotline numbers on emergency trigger", "Read", "Emergency"],
        ["FR40", "Log all emergency button activations", "Create", "Audit"],
        ["FR41", "Add mental health resources via admin panel", "Create", "Content"],
        ["FR42", "Validate resource URLs before publishing", "Create", "Content"],
        ["FR43", "Track and aggregate resource view counts", "Read", "Analytics"],
        ["FR44", "Generate monthly platform analytics reports", "Read", "Analytics"],
        ["FR45", "Execute automated daily database backups at midnight", "Create", "System"],
        ["FR46", "Restore database from backup within 1-hour window", "Read", "System"],
        ["FR47", "Log system errors and exceptions automatically", "Create", "Monitoring"],
        ["FR48", "Dispatch instant alerts for critical system errors", "Create", "Monitoring"],
        ["FR49", "Select preferred interface language", "Update", "UI"],
        ["FR50", "Load main dashboard within 3 seconds", "Read", "Performance"],
        ["FR51", "Upload clinical documents and worksheets", "Create", "Clinical"],
        ["FR52", "Validate document uploads (PDF format only)", "Create", "Clinical"],
        ["FR53", "Download personal medical records securely", "Read", "Compliance"],
        ["FR54", "Revoke data sharing consent instantaneously", "Update", "Security"],
        ["FR55", "Restrict all access upon consent revocation", "Update", "Security"],
        ["FR56", "Log all modifications to user consent settings", "Create", "Audit"],
        ["FR57", "Integrate external SMS API for notifications", "Create", "Integration"],
        ["FR58", "Retry failed API calls up to 3 times", "Update", "Integration"],
        ["FR59", "Queue notifications during service downtimes", "Create", "Integration"],
        ["FR60", "Deactivate user accounts via admin panel", "Update", "Admin"],
        ["FR61", "Track active session duration per user", "Read", "Analytics"],
        ["FR62", "Auto-logout after 15 minutes of inactivity", "Update", "Security"],
        ["FR63", "Validate inputs to prevent SQL injection", "Create", "Security"],
        ["FR64", "Generate complete audit trails for DB transactions", "Create", "Audit"],
        ["FR65", "Assign and enforce RBAC", "Update", "Security"],
        ["FR66", "Deny and log unauthorized API requests", "Read", "Security"],
        ["FR67", "Synchronize appointments with Google Calendar", "Create", "Integration"],
        ["FR68", "Dispatch follow-up for missed appointments", "Create", "Notification"],
        ["FR69", "Rebook from missed session alerts", "Create", "Appointment"],
        ["FR70", "Track clinician workload metrics", "Read", "Analytics"],
        ["FR71", "Generate regulatory compliance reports", "Read", "Compliance"],
        ["FR72", "Log complete history of data modifications", "Create", "Audit"],
        ["FR73", "Restrict file uploads exceeding 5MB", "Create", "System"],
        ["FR74", "Export reports to CSV", "Read", "Admin"],
        ["FR75", "Maintain system monthly", "Read", "System"],
    ]
    add_styled_table(doc, ["FR#", "Requirement", "CRUD", "Module"], fr_data, col_widths=[0.5, 3.8, 0.6, 1.0])

    # 2.2 Non-Functional Requirements
    add_heading_styled(doc, "2.2 Non-Functional Requirements", level=2)

    nfr_categories = [
        ("Performance (NFRP01–NFRP04)", [
            ["NFRP01", "Page load time ≤ 3 seconds", "Performance"],
            ["NFRP02", "API response 95th percentile within 500ms", "Performance"],
            ["NFRP03", "Support 50 concurrent users minimum", "Scalability"],
            ["NFRP04", "Database queries complete within 200ms", "Performance"],
        ]),
        ("Security (NFRS01–NFRS06)", [
            ["NFRS01", "Password hashing using bcrypt with 10 salt rounds", "Security"],
            ["NFRS02", "JWT tokens with 1-hour expiry", "Security"],
            ["NFRS03", "HTTPS/TLS 1.2+ for all communications", "Security"],
            ["NFRS04", "AES-256 encryption for data at rest", "Security"],
            ["NFRS05", "SQL injection and XSS prevention on all inputs", "Security"],
            ["NFRS06", "Rate limiting: 100 requests/IP/minute (HTTP 429)", "Security"],
        ]),
        ("Usability (NFRU01–NFRU04)", [
            ["NFRU01", "Intuitive interface with ≤ 5 minute onboarding", "Usability"],
            ["NFRU02", "Responsive design: 320px – 1920px viewport", "Usability"],
            ["NFRU03", "WCAG 2.1 Level AA compliance", "Accessibility"],
            ["NFRU04", "Descriptive and actionable error messages", "Usability"],
        ]),
        ("Reliability (NFRR01–NFRR03)", [
            ["NFRR01", "99% system uptime", "Reliability"],
            ["NFRR02", "Graceful degradation for non-critical features", "Reliability"],
            ["NFRR03", "Daily automated backups with 7-day retention", "Reliability"],
        ]),
        ("Maintainability (NFRM01–NFRM03)", [
            ["NFRM01", "ESLint enforced code standards", "Maintainability"],
            ["NFRM02", "Modular architecture with clear separation of concerns", "Maintainability"],
            ["NFRM03", "Scalable design supporting 10x user growth", "Scalability"],
        ]),
        ("Privacy & Compliance (NFRC01–NFRC04)", [
            ["NFRC01", "Privacy by design principles throughout", "Privacy"],
            ["NFRC02", "Granular consent management", "Compliance"],
            ["NFRC03", "Data anonymisation for analytics", "Privacy"],
            ["NFRC04", "Right to deletion within 30 days", "Compliance"],
        ]),
    ]
    for cat_title, rows in nfr_categories:
        add_body_text(doc, cat_title, bold=True, size=11)
        add_styled_table(doc, ["ID", "Requirement", "Category"], rows, col_widths=[1.0, 4.2, 1.2])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 3. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "3. System Architecture", level=1)

    add_heading_styled(doc, "3.1 Software and Hardware Architecture", level=2)
    add_body_text(doc, "The Mental Health Support Platform follows a three-tier architecture pattern comprising a Presentation Tier (Client), Application Tier (Server), and Data Tier (Storage). This architecture ensures separation of concerns, scalability, and maintainability.")

    svg_path = os.path.join(BASE_DIR, "3. System Architecture_diagram_1.svg")
    add_svg_image(doc, svg_path, "Figure 3.1: Three-Tier System Architecture Diagram", 6.3)

    add_body_text(doc, "Architecture Explanation:", bold=True)
    add_body_text(doc, "Presentation Tier (Client): Built with React.js and Vite using TypeScript. The client layer includes a Single Page Application (SPA) with component-based routing, React Context and React Query for state management, Socket.io client for real-time chat and notifications, Chart.js for mood analytics dashboards, and Axios HTTP client with JWT interceptor for secure API communication.")
    add_body_text(doc, "Application Tier (Server): Powered by Express.js running on Node.js. The server tier includes a REST API server with middleware chain and error handling, an authentication module (JWT + OTP + bcrypt hashing + RBAC), business logic for appointments, mood tracking, and clinical records, Socket.io server for real-time chat engine with AES-256 encryption and presence tracking, and external API integrations with OpenAI (chatbot), Twilio (SMS), and SendGrid (email).")
    add_body_text(doc, "Data Tier (Storage): Uses PostgreSQL as the primary relational database storing users, appointments, clinical records, messages, and mood entries. Additional storage components include file storage for profile images, clinical documents, and backup archives; audit logs for access, error, and change history tracking; and cache/sessions for JWT token store, session data, and API response caching.")
    add_body_text(doc, "Communication: The client communicates with the server over HTTPS and WSS (WebSocket Secure) protocols. Server-to-database communication uses SQL queries through an ORM layer.")

    # 3.2 Tech Stack
    add_heading_styled(doc, "3.2 Tech Stack (Detailed)", level=2)
    tech_data = [
        ["Frontend Framework", "React.js 18+ with Vite", "Component-based SPA with TypeScript support"],
        ["State Management", "React Context + React Query", "Server state caching and local state management"],
        ["UI Styling", "Tailwind CSS / Material UI", "Responsive design, utility-first CSS"],
        ["Charts & Visualization", "Chart.js / Recharts", "Mood trend visualization, analytics dashboards"],
        ["Real-time Communication", "Socket.io (Client)", "WebSocket-based chat and notifications"],
        ["HTTP Client", "Axios", "API communication with JWT interceptor"],
        ["Backend Runtime", "Node.js 18+", "Server-side JavaScript runtime"],
        ["API Framework", "Express.js", "RESTful API with middleware architecture"],
        ["Authentication", "JWT + bcrypt + OTP", "Stateless auth, password hashing, 2FA"],
        ["Chat Encryption", "AES-256", "End-to-end message encryption"],
        ["Database", "PostgreSQL 15+", "Primary relational database (ACID compliant)"],
        ["ORM", "Prisma / Sequelize", "Database abstraction and migration management"],
        ["AI/NLP", "OpenAI API (GPT-3.5/4)", "Chatbot NLP processing"],
        ["SMS Gateway", "Twilio", "SMS notifications and OTP delivery"],
        ["Email Service", "SendGrid", "Transactional email notifications"],
        ["Calendar Integration", "Google Calendar API", "Appointment synchronization"],
        ["Testing (Unit)", "Jest", "Unit test framework for Node.js and React"],
        ["Testing (E2E)", "Cypress", "End-to-end browser testing"],
        ["Version Control", "Git / GitHub", "Source code management and collaboration"],
        ["CI/CD", "GitHub Actions", "Automated testing and deployment pipeline"],
        ["Cloud Hosting", "Render / AWS", "Application and database hosting"],
        ["Wireframing", "Balsamiq Cloud", "Low-fidelity wireframe design"],
    ]
    add_styled_table(doc, ["Component", "Technology", "Purpose"], tech_data, col_widths=[1.6, 2.0, 3.0])

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 4. SYSTEM DESIGN
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "4. System Design", level=1)
    add_heading_styled(doc, "4.1 User Interface Design / Wireframes", level=2)
    add_body_text(doc, "The following wireframes were designed using Balsamiq Cloud to represent the key user interface screens of the Mental Health Support Platform. Each wireframe demonstrates the layout, navigation, and key interactive elements.")

    wireframe_files = [
        ("4.1 User interface Design Wireframes_diagram_1.svg", "Figure 4.1: Login / Registration Page"),
        ("4.1 User interface Design Wireframes_diagram_2.svg", "Figure 4.2: User Dashboard"),
        ("4.1 User interface Design Wireframes_diagram_3.svg", "Figure 4.3: Mood Tracker Page"),
        ("4.1 User interface Design Wireframes_diagram_4.svg", "Figure 4.4: Secure Messaging Interface"),
        ("4.1 User interface Design Wireframes_diagram_5.svg", "Figure 4.5: Appointments Page"),
        ("4.1 User interface Design Wireframes_diagram_6.svg", "Figure 4.6: Admin Dashboard Panel"),
    ]
    for fname, caption in wireframe_files:
        svg_path = os.path.join(BASE_DIR, fname)
        add_svg_image(doc, svg_path, caption, 5.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 5. USE CASE DIAGRAMS
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "5. Use Case Diagrams", level=1)
    add_body_text(doc, "The use case diagrams illustrate the interactions between the three primary actors (Patient/User, Clinician/Doctor, and Administrator) and the Mental Health Support Platform. Each use case is mapped to specific functional requirements.")

    uc_files = [
        ("Patient  User — Use Case Diagram.svg", "Figure 5.1: Patient / User — Use Case Diagram"),
        ("Clinician  Doctor — Use Case Diagram.svg", "Figure 5.2: Clinician / Doctor — Use Case Diagram"),
        ("Administrator — Use Case Diagram.svg", "Figure 5.3: Administrator — Use Case Diagram"),
    ]
    for fname, caption in uc_files:
        svg_path = os.path.join(BASE_DIR, fname)
        add_svg_image(doc, svg_path, caption, 6.0)

    # Use case descriptions
    add_body_text(doc, "Use Case Summary — Patient/User:", bold=True)
    patient_ucs = [
        "Register Account (FR1), Login/Logout (FR8), Verify OTP (FR2), Reset Password (FR5)",
        "Update Profile (FR6), Upload Profile Image (FR7)",
        "Submit Mood Entry (FR27), View Mood Report (FR29)",
        "Book Appointment (FR9), Reschedule Appointment (FR11)",
        "Send Chat Message (FR23), Use AI Chatbot (FR32)",
        "Access Resources (FR41), Trigger Emergency (FR38)",
        "Manage Consent (FR54), View Dashboard (FR50), Download Records (FR53), Select Language (FR49)",
    ]
    for uc in patient_ucs:
        add_bullet(doc, uc)

    add_body_text(doc, "Use Case Summary — Clinician:", bold=True)
    clinician_ucs = [
        "Login/Logout (FR8), View Patient History (FR18), Write Session Notes (FR16)",
        "Approve Appointments (FR14), Chat with Patient (FR24), View Schedule (FR67)",
        "Request Medical Records (FR21), Reassign Patients (FR20)",
        "Upload Clinical Docs (FR51), View Workload Metrics (FR70), Access Patient Mood Data (FR29)",
    ]
    for uc in clinician_ucs:
        add_bullet(doc, uc)

    add_body_text(doc, "Use Case Summary — Administrator:", bold=True)
    admin_ucs = [
        "Manage Users (FR60), Assign RBAC Roles (FR65), Manage Resources (FR41), Validate URLs (FR42)",
        "View Analytics Dashboard (FR44), Generate Compliance Reports (FR71)",
        "Review Chatbot Logs (FR35), Monitor System Health (FR47)",
        "Create/Restore Backups (FR45–46), Export Reports to CSV (FR74)",
        "Deactivate Accounts (FR60), View Audit Trails (FR64)",
    ]
    for uc in admin_ucs:
        add_bullet(doc, uc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 6. CONTEXT DIAGRAM
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "6. Context Diagrams (Gane & Sarson Method)", level=1)
    add_body_text(doc, "The Level 0 Context Diagram provides a high-level view of the Mental Health Support Platform and its interactions with external entities. This diagram uses the Gane & Sarson notation where processes are represented as rounded rectangles and external entities as rectangular boxes.")

    svg_path = os.path.join(BASE_DIR, "Level 0 — Context Diagram (Gane & Sarson).svg")
    add_svg_image(doc, svg_path, "Figure 6.1: Level 0 — Context Diagram (Gane & Sarson)", 6.3)

    add_body_text(doc, "External Entities:", bold=True)
    entities = [
        "E1: Patient/User — Primary system user who registers, tracks mood, books appointments, and accesses resources",
        "E2: Clinician — Healthcare professional who manages appointments, writes session notes, and views patient records",
        "E3: Administrator — System administrator who manages users, resources, analytics, and system configuration",
        "E4: SMS Gateway (Twilio) — External service for sending SMS notifications and OTP codes",
        "E5: Google Calendar — External calendar service for appointment synchronization",
        "E6: Email Service (SendGrid) — External service for transactional email notifications",
        "E7: Crisis Hotline — External crisis support service providing hotline data for emergency scenarios",
    ]
    for e in entities:
        add_bullet(doc, e)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 7. DATA FLOW DIAGRAMS
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "7. Data Flow Diagrams (Gane & Sarson Method)", level=1)
    add_body_text(doc, "The Data Flow Diagrams (DFDs) decompose the system into progressively detailed levels, showing how data flows between processes, external entities, and data stores. All diagrams follow the Gane & Sarson notation.")

    # 7.1 DFD Level 1
    add_heading_styled(doc, "7.1 DFD Level 1", level=2)
    add_body_text(doc, "The Level 1 DFD decomposes the central system process into 15 major sub-processes, each mapped to specific functional requirements. The diagram shows data flows between external entities (E1–E7), processes (1.0–15.0), and 15 data stores (D1–D15).")

    svg_path = os.path.join(BASE_DIR, "Level 1 — System Data Flow Diagram.svg")
    add_svg_image(doc, svg_path, "Figure 7.1: Level 1 — System Data Flow Diagram", 6.3)

    # Process descriptions
    processes_l1 = [
        ["1.0", "Authentication & Account Management", "FR1–5, FR8"],
        ["2.0", "Profile Management", "FR6–7"],
        ["3.0", "Appointment Management", "FR9–15, FR67–69"],
        ["4.0", "Clinical Records Management", "FR16–22, FR51–53"],
        ["5.0", "Messaging System", "FR23–26"],
        ["6.0", "Mood Tracking & Analysis", "FR27–31"],
        ["7.0", "AI Chatbot System", "FR32–36"],
        ["8.0", "Emergency Support", "FR38–40"],
        ["9.0", "Notification Engine", "FR13, FR37, FR57–59, FR68"],
        ["10.0", "Content & Resource Management", "FR41–43"],
        ["11.0", "Analytics & Reporting", "FR44, FR61, FR70–72, FR74"],
        ["12.0", "Administration & Access Control", "FR60, FR62–66"],
        ["13.0", "System Maintenance & Backup", "FR45–48, FR73, FR75"],
        ["14.0", "Privacy & Compliance", "FR54–56"],
        ["15.0", "UI & Localization", "FR49–50"],
    ]
    add_styled_table(doc, ["Process", "Name", "FRs Covered"], processes_l1, col_widths=[0.8, 2.8, 2.5])

    # Data Stores
    doc.add_paragraph()
    add_body_text(doc, "Data Stores:", bold=True)
    data_stores = [
        ["D1", "Users & Credentials"], ["D2", "User Profiles"], ["D3", "Appointments"],
        ["D4", "Clinical Records"], ["D5", "Chat Messages"], ["D6", "Mood Entries"],
        ["D7", "Chatbot Transcripts"], ["D8", "Emergency Logs"], ["D9", "Notification Queue"],
        ["D10", "Resource Library"], ["D11", "Analytics Data"], ["D12", "Audit Logs"],
        ["D13", "System Backups"], ["D14", "Consent Records"], ["D15", "Session Data"],
    ]
    add_styled_table(doc, ["Store ID", "Name"], data_stores, col_widths=[1.0, 4.0])

    doc.add_page_break()

    # 7.2 DFD Level 2
    add_heading_styled(doc, "7.2 DFD Level 2", level=2)
    add_body_text(doc, "The Level 2 DFDs decompose each Level 1 process into more detailed sub-processes. Each Level 2 diagram shows the internal data flows, decision logic, and data store interactions within its parent process.")

    level2_files = [
        ("Level 2.1 — Authentication & Account Management.svg", "Figure 7.2.1: Level 2.1 — Authentication & Account Management", "Decomposes Process 1.0 into 7 sub-processes: Register User (FR1), Verify OTP (FR2), Lock Account (FR3), Log Login Activity (FR4), Password Reset (FR5), Login/Logout (FR8), and Delete & Anonymize (FR8d)."),
        ("Level 2.2 — Profile Management.svg", "Figure 7.2.2: Level 2.2 — Profile Management", "Decomposes Process 2.0 into 2 sub-processes: Update Profile Fields (FR6) and Upload Profile Image with < 2MB validation (FR7)."),
        ("Level 2.3 — Appointment Management.svg", "Figure 7.2.3: Level 2.3 — Appointment Management", "Decomposes Process 3.0 into 9 sub-processes covering scheduling (FR9), slot availability (FR10), rescheduling (FR11), missed marking (FR12), approval pending (FR14), auto-cancel (FR15), Google Calendar sync (FR67), follow-up notifications (FR68), and rebooking from alerts (FR69)."),
        ("Level 2.4 — Clinical Records Management.svg", "Figure 7.2.4: Level 2.4 — Clinical Records Management", "Decomposes Process 4.0 into 10 sub-processes managing session notes (FR16), consent checking (FR17), patient history viewing (FR18), record retention (FR19), patient reassignment (FR20), medical record requests (FR21), access logging (FR22), clinical document upload (FR51), format validation (FR52), and record download (FR53)."),
        ("Level 2.5 — Messaging System.svg", "Figure 7.2.5: Level 2.5 — Messaging System", "Decomposes Process 5.0 into 4 sub-processes: message encryption with AES-256 (FR23), message delivery within 2 seconds (FR24), chat transcript storage (FR25), and soft delete functionality (FR26)."),
        ("Level 2.6 — Mood Tracking & Analysis.svg", "Figure 7.2.6: Level 2.6 — Mood Tracking & Analysis", "Decomposes Process 6.0 into 5 sub-processes: mood entry submission (FR27), duplicate prevention within 1-hour window (FR28), weekly mood report generation (FR29), risk alert triggering for consecutive low scores (FR30), and resource recommendation based on patterns (FR31)."),
        ("Level 2.7 — AI Chatbot System.svg", "Figure 7.2.7: Level 2.7 — AI Chatbot System", "Decomposes Process 7.0 into 5 sub-processes: query processing within 3 seconds (FR32), human escalation when confidence < 70% (FR33), transcript storage (FR34), admin log review (FR35), and inactive user detection after 7 days (FR36)."),
        ("Level 2.8 — Emergency Support.svg", "Figure 7.2.8: Level 2.8 — Emergency Support", "Decomposes Process 8.0 into 3 sub-processes: emergency button activation with < 1 second response (FR38), crisis hotline display (FR39), and emergency activation logging (FR40)."),
        ("Level 2.9 — Notification Engine.svg", "Figure 7.2.9: Level 2.9 — Notification Engine", "Decomposes Process 9.0 into 6 sub-processes covering booking confirmation (FR13), wellness check dispatch (FR37), SMS/email integration (FR57), failed call retry up to 3 times (FR58), downtime queuing (FR59), and missed appointment follow-up (FR68)."),
        ("Level 2.10 — Content & Resource Management.svg", "Figure 7.2.10: Level 2.10 — Content & Resource Management", "Decomposes Process 10.0 into 3 sub-processes: resource addition via admin panel (FR41), URL validation before publishing (FR42), and view count tracking (FR43)."),
        ("Level 2.11 — Analytics & Reporting.svg", "Figure 7.2.11: Level 2.11 — Analytics & Reporting", "Decomposes Process 11.0 into 6 sub-processes: monthly report generation (FR44), session duration tracking (FR61), clinician workload evaluation (FR70), compliance reporting (FR71), data modification logging (FR72), and CSV export (FR74)."),
        ("Level 2.12 — Administration & Access Control.svg", "Figure 7.2.12: Level 2.12 — Administration & Access Control", "Decomposes Process 12.0 into 6 sub-processes: account deactivation (FR60), auto-logout after 15 minutes (FR62), input validation for SQL injection prevention (FR63), audit trail generation (FR64), RBAC assignment (FR65), and unauthorized API request denial (FR66)."),
        ("Level 2.13 — System Maintenance & Backup.svg", "Figure 7.2.13: Level 2.13 — System Maintenance & Backup", "Decomposes Process 13.0 into 6 sub-processes: daily midnight backup (FR45), backup restoration (FR46), system error logging (FR47), critical admin alerts (FR48), file upload limit enforcement at 5MB (FR73), and monthly maintenance (FR75)."),
        ("Level 2.14 — Privacy & Compliance.svg", "Figure 7.2.14: Level 2.14 — Privacy & Compliance", "Decomposes Process 14.0 into 3 sub-processes: instant consent revocation (FR54), access restriction on revocation (FR55), and consent change logging (FR56)."),
        ("Level 2.15 — UI & Localization.svg", "Figure 7.2.15: Level 2.15 — UI & Localization", "Decomposes Process 15.0 into 2 sub-processes: interface language selection (FR49) and dashboard loading within 3 seconds (FR50)."),
    ]

    for fname, caption, description in level2_files:
        svg_path = os.path.join(BASE_DIR, fname)
        add_body_text(doc, description, size=10)
        add_svg_image(doc, svg_path, caption, 6.0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 8. SEQUENCE DIAGRAMS
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "8. Sequence Diagrams", level=1)
    add_body_text(doc, "The sequence diagrams illustrate the dynamic behavior of the system by showing the order of messages exchanged between objects (actors, frontend, API server, database, and external services) during key system interactions.")

    sd_files = [
        ("SD1 User Registration & Login.svg", "Figure 8.1: SD1 — User Registration & Login",
         "This sequence diagram shows the complete user registration flow including data submission, database record creation, OTP email verification, OTP code verification, account verification update, JWT token generation, and dashboard rendering. The flow involves Patient, Frontend, API Server, and Database participants."),
        ("SD2 Appointment Booking Flow.svg", "Figure 8.2: SD2 — Appointment Booking Flow",
         "This diagram illustrates the appointment booking process from clinician/date selection through slot availability checking, booking confirmation, database insertion, Google Calendar synchronization, and email notification dispatch. It involves Patient, Frontend, API Server, Database, and Google Calendar participants."),
        ("SD3 Mood Entry & Risk Alert.svg", "Figure 8.3: SD3 — Mood Entry & Risk Alert",
         "Shows the mood submission flow with duplicate checking (24-hour window), mood entry insertion, consecutive low-score detection (< 2 for 3 days), risk alert triggering, notification insertion, and dashboard update. Demonstrates the automated risk detection mechanism."),
        ("SD4 AI Chatbot Interaction.svg", "Figure 8.4: SD4 — AI Chatbot Interaction",
         "Illustrates the AI chatbot conversation flow: message submission, API server processing, NLP query to AI engine, confidence scoring, response delivery (< 3 seconds), transcript storage, and automatic escalation to human clinician when confidence drops below 70%."),
        ("SD5 Emergency Button Activation.svg", "Figure 8.5: SD5 — Emergency Button Activation",
         "Shows the critical emergency flow: button press, emergency activation API call, emergency log creation, crisis hotline data fetch from external service, crisis information display (< 1 second response), and audit trail insertion. Involves Patient, Frontend, API Server, Database, and Crisis Service."),
        ("SD6 Secure Messaging.svg", "Figure 8.6: SD6 — Secure Messaging",
         "Demonstrates the secure messaging flow: message composition, AES-256 encryption at the server, encrypted message storage, WebSocket real-time delivery (< 2 seconds), delivery confirmation, recipient notification, and soft delete functionality."),
    ]

    for fname, caption, description in sd_files:
        svg_path = os.path.join(BASE_DIR, fname)
        add_body_text(doc, description, size=10)
        add_svg_image(doc, svg_path, caption, 5.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 9. ERD & DATA DICTIONARY
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "9. Entity Relationship Diagram & Data Dictionary", level=1)
    add_body_text(doc, "The Entity Relationship Diagram (ERD) represents the logical data model of the Mental Health Support Platform database. The diagram follows cross-reference modulation principles with clearly defined primary keys (PK), foreign keys (FK), and cardinality relationships.")

    svg_path = os.path.join(BASE_DIR, "Entity Relationship Diagram — Mental Health Support Platform.svg")
    add_svg_image(doc, svg_path, "Figure 9.1: Entity Relationship Diagram — Mental Health Support Platform", 6.3)

    add_heading_styled(doc, "9.1 Data Dictionary", level=2)
    add_body_text(doc, "The following tables define the schema for each entity in the database:")

    # Data dictionary entries
    dd_tables = [
        ("Users", [
            ["user_id", "UUID", "PK", "Unique user identifier"],
            ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "User email address"],
            ["password_hash", "TEXT", "NOT NULL", "bcrypt hashed password (10 salt rounds)"],
            ["role", "ENUM", "NOT NULL", "patient | clinician | admin"],
            ["verified", "BOOLEAN", "DEFAULT false", "Email/OTP verification status"],
            ["locked", "BOOLEAN", "DEFAULT false", "Account lock after 5 failed attempts"],
            ["mfa_enabled", "BOOLEAN", "DEFAULT false", "Multi-factor authentication flag"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Account creation timestamp"],
            ["last_login", "TIMESTAMP", "NULLABLE", "Most recent login timestamp"],
        ]),
        ("Profiles", [
            ["profile_id", "UUID", "PK", "Unique profile identifier"],
            ["user_id", "UUID", "FK → Users", "References Users table"],
            ["display_name", "VARCHAR(100)", "NOT NULL", "User display name"],
            ["avatar_url", "TEXT", "NULLABLE", "Profile image URL (< 2MB)"],
            ["language", "VARCHAR(10)", "DEFAULT 'en'", "Interface language preference"],
            ["updated_at", "TIMESTAMP", "DEFAULT NOW()", "Last profile update"],
        ]),
        ("Appointments", [
            ["appt_id", "UUID", "PK", "Unique appointment identifier"],
            ["patient_id", "UUID", "FK → Users", "Patient reference"],
            ["clinician_id", "UUID", "FK → Users", "Clinician reference"],
            ["date_time", "TIMESTAMP", "NOT NULL", "Appointment date and time"],
            ["status", "ENUM", "NOT NULL", "pending | confirmed | missed | cancelled"],
            ["gcal_event_id", "TEXT", "NULLABLE", "Google Calendar event ID"],
            ["missed", "BOOLEAN", "DEFAULT false", "Missed after 15-min delay"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Booking creation time"],
            ["auto_cancelled", "BOOLEAN", "DEFAULT false", "Auto-cancelled after 30 min"],
        ]),
        ("ClinicalRecords", [
            ["record_id", "UUID", "PK", "Unique record identifier"],
            ["patient_id", "UUID", "FK → Users", "Patient reference"],
            ["clinician_id", "UUID", "FK → Users", "Clinician reference"],
            ["session_notes", "TEXT", "NOT NULL", "Clinical session notes"],
            ["record_type", "ENUM", "NOT NULL", "consultation | follow-up | emergency"],
            ["retention_date", "DATE", "NULLABLE", "Data retention deadline"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Record creation time"],
        ]),
        ("ChatMessages", [
            ["message_id", "UUID", "PK", "Unique message identifier"],
            ["sender_id", "UUID", "FK → Users", "Message sender"],
            ["receiver_id", "UUID", "FK → Users", "Message receiver"],
            ["encrypted_body", "TEXT", "NOT NULL", "AES-256 encrypted message body"],
            ["delivered", "BOOLEAN", "DEFAULT false", "Delivery confirmation"],
            ["soft_deleted", "BOOLEAN", "DEFAULT false", "Soft delete flag (not purged)"],
        ]),
        ("MoodEntries", [
            ["mood_id", "UUID", "PK", "Unique mood entry identifier"],
            ["user_id", "UUID", "FK → Users", "User reference"],
            ["score", "INT", "CHECK (1-5)", "Mood score (1=very low, 5=excellent)"],
            ["notes", "TEXT", "NULLABLE", "Optional mood notes"],
            ["entry_date", "DATE", "NOT NULL", "Date of mood entry"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Entry creation time"],
        ]),
        ("ChatbotTranscripts", [
            ["transcript_id", "UUID", "PK", "Unique transcript identifier"],
            ["user_id", "UUID", "FK → Users", "User reference"],
            ["query", "TEXT", "NOT NULL", "User query text"],
            ["response", "TEXT", "NOT NULL", "AI-generated response"],
            ["confidence", "FLOAT", "NOT NULL", "AI confidence score (0.0–1.0)"],
            ["escalated", "BOOLEAN", "DEFAULT false", "Escalated to human (< 70%)"],
        ]),
        ("Notifications", [
            ["notif_id", "UUID", "PK", "Unique notification identifier"],
            ["user_id", "UUID", "FK → Users", "Recipient user reference"],
            ["type", "ENUM", "NOT NULL", "booking | wellness | alert | follow_up"],
            ["channel", "ENUM", "NOT NULL", "email | sms | in_app"],
            ["payload", "JSON", "NOT NULL", "Notification content payload"],
            ["status", "ENUM", "DEFAULT 'pending'", "pending | sent | failed | queued"],
            ["retry_count", "INT", "DEFAULT 0", "API retry count (max 3)"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Notification creation time"],
        ]),
        ("EmergencyLogs", [
            ["log_id", "UUID", "PK", "Unique emergency log identifier"],
            ["user_id", "UUID", "FK → Users", "User who activated emergency"],
            ["activated_at", "TIMESTAMP", "DEFAULT NOW()", "Activation timestamp"],
            ["hotline_data", "JSON", "NULLABLE", "Crisis hotline information"],
            ["resolved", "BOOLEAN", "DEFAULT false", "Resolution status"],
        ]),
        ("ConsentRecords", [
            ["consent_id", "UUID", "PK", "Unique consent identifier"],
            ["user_id", "UUID", "FK → Users", "User reference"],
            ["consent_type", "ENUM", "NOT NULL", "data_sharing | analytics | marketing"],
            ["granted", "BOOLEAN", "NOT NULL", "Consent granted or revoked"],
            ["changed_at", "TIMESTAMP", "DEFAULT NOW()", "Last consent change time"],
        ]),
        ("AuditLogs", [
            ["audit_id", "UUID", "PK", "Unique audit entry identifier"],
            ["user_id", "UUID", "FK → Users", "User who performed action"],
            ["action", "VARCHAR(100)", "NOT NULL", "Action performed"],
            ["entity_type", "VARCHAR(50)", "NOT NULL", "Affected entity type"],
            ["entity_id", "UUID", "NULLABLE", "Affected entity ID"],
            ["ip_address", "INET", "NULLABLE", "Client IP address"],
            ["details", "JSON", "NULLABLE", "Additional action details"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "Audit entry time"],
        ]),
    ]

    for table_name, columns in dd_tables:
        add_body_text(doc, f"Table: {table_name}", bold=True, size=11)
        add_styled_table(doc, ["Column", "Type", "Constraints", "Description"], columns, col_widths=[1.4, 1.3, 1.5, 2.4])
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 10. CLASS DIAGRAM
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "10. Class Diagram", level=1)
    add_body_text(doc, "The UML Class Diagram represents the object-oriented model of the Mental Health Support Platform. Each class defines attributes (properties) and methods (operations), with relationships showing associations, dependencies, and multiplicities between classes.")

    svg_path = os.path.join(BASE_DIR, "UML Class Diagram — Mental Health Support Platform.svg")
    add_svg_image(doc, svg_path, "Figure 10.1: UML Class Diagram — Mental Health Support Platform", 6.3)

    add_body_text(doc, "The system comprises 12 primary classes:", bold=True)
    classes_desc = [
        "User — Core authentication entity with registration, login, OTP verification, password reset, and account deletion methods",
        "Profile — User profile management with display name, avatar upload, and language preferences",
        "Appointment — Scheduling system with conflict detection, slot checking, rescheduling, approval, auto-cancel, and missed marking",
        "ClinicalRecord — Medical records management with consent checking, session notes, document upload/download, and patient history viewing",
        "ChatMessage — Encrypted messaging (AES-256) with send, deliver, soft-delete operations",
        "MoodEntry — Daily mood tracking with duplicate prevention, weekly reporting, risk alerting, and resource recommendation",
        "Chatbot — AI-powered conversation with NLP query processing, human escalation, transcript storage, and inactive user detection",
        "Emergency — Emergency support activation with hotline fetching and activation logging",
        "Notification — Multi-channel notification system (email, SMS, in-app) with retry and downtime queuing",
        "Resource — Mental health resource library with URL validation, view tracking, and content management",
        "Analytics — Reporting engine with monthly reports, clinician workload metrics, compliance reports, and CSV export",
        "Admin — Administrative controls with user management, RBAC assignment, account deactivation, audit trail viewing, and backup management",
    ]
    for c in classes_desc:
        add_bullet(doc, c)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 11. DPIA
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "11. Data Protection Impact Assessment (DPIA)", level=1)
    add_body_text(doc, "The Data Protection Impact Assessment identifies and mitigates privacy risks associated with the processing of personal and sensitive health data within the Mental Health Support Platform.")

    add_heading_styled(doc, "11.1 Overview of Data Processing", level=2)
    add_body_text(doc, "The platform processes several categories of personal data:")
    data_cats = [
        "Account Data — Email addresses, hashed passwords, login timestamps, IP addresses",
        "Profile Data — Display names, avatar images, language preferences",
        "Health Data — Mood scores, mood notes, clinical session notes, medical records",
        "Communication Data — Encrypted chat messages, chatbot transcripts",
        "Behavioral Data — Session duration, platform usage patterns, resource view counts",
        "Emergency Data — Emergency activation logs, crisis interaction records",
        "Consent Data — Data sharing preferences, consent change history",
    ]
    for c in data_cats:
        add_bullet(doc, c)

    add_heading_styled(doc, "11.2 Risk Assessment", level=2)
    risk_data = [
        ["Unauthorized access to health data", "High", "RBAC, JWT auth, consent-based access (FR17, FR65)", "Low"],
        ["Data breach of personal information", "High", "AES-256 encryption, bcrypt hashing, HTTPS (FR23, NFRS03)", "Low"],
        ["Unintended data retention", "Medium", "Right to deletion (FR8d), retention policies (FR19)", "Low"],
        ["Unauthorized data sharing", "High", "Consent management (FR54–55), immediate revocation", "Low"],
        ["Session hijacking", "Medium", "JWT 1hr expiry, auto-logout 15min (FR62, NFRS02)", "Low"],
        ["SQL injection / XSS attacks", "High", "Input validation (FR63), parameterized queries (NFRS05)", "Low"],
        ["Loss of data", "Medium", "Daily automated backups (FR45), restore within 1hr (FR46)", "Low"],
        ["Excessive API access", "Medium", "Rate limiting 100 req/IP/min (NFRS06)", "Low"],
        ["Emergency data misuse", "Medium", "Audit logging (FR40), access logging (FR22)", "Low"],
        ["Non-compliance with privacy laws", "High", "GDPR alignment, consent management (NFRC01–04)", "Low"],
    ]
    add_styled_table(doc, ["Risk", "Severity", "Mitigation Measures", "Residual"], risk_data, col_widths=[1.8, 0.7, 3.3, 0.7])

    add_heading_styled(doc, "11.3 Privacy Controls Summary", level=2)
    privacy_controls = [
        ["Encryption at Rest", "AES-256 for all sensitive data including messages and clinical records"],
        ["Encryption in Transit", "HTTPS/TLS 1.2+ for all client-server communications"],
        ["Password Security", "bcrypt hashing with 10 salt rounds; no plaintext storage"],
        ["Authentication", "JWT tokens with 1-hour expiry; OTP verification; multi-factor authentication"],
        ["Access Control", "Role-Based Access Control (RBAC) with Patient, Clinician, and Admin roles"],
        ["Consent Management", "Granular consent with instant revocation and immediate access restriction"],
        ["Audit Trails", "Complete logging of all database transactions, record access, and consent changes"],
        ["Data Minimization", "Only essential data collected; data anonymisation for analytics"],
        ["Right to Deletion", "User account deletion with data anonymisation within 24 hours"],
        ["Backup & Recovery", "Daily automated backups at midnight with 7-day retention; 1-hour restore SLA"],
    ]
    add_styled_table(doc, ["Control", "Implementation"], privacy_controls, col_widths=[1.8, 4.8])

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 12. TEST AND IMPLEMENTATION PLAN
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "12. Test and Implementation Plan", level=1)

    add_heading_styled(doc, "12.1 Testing Strategy", level=2)
    add_body_text(doc, "The testing strategy follows a multi-layered approach to ensure comprehensive quality assurance across all system components:")

    test_layers = [
        ["Unit Testing", "Jest", "Individual functions, API endpoints, React components", "≥ 80% code coverage"],
        ["Integration Testing", "Jest + Supertest", "API endpoint chains, database interactions, auth flows", "All critical paths"],
        ["End-to-End Testing", "Cypress", "Complete user workflows from UI to database", "Core user journeys"],
        ["Security Testing", "OWASP ZAP / Manual", "SQL injection, XSS, auth bypass, rate limiting", "Zero critical vulnerabilities"],
        ["Performance Testing", "Artillery / k6", "Page load < 3s, API response < 500ms, 50 concurrent users", "Meet all NFR targets"],
        ["User Acceptance Testing", "Manual", "Stakeholder validation of all features", "All FRs verified"],
    ]
    add_styled_table(doc, ["Test Type", "Tool", "Scope", "Acceptance Criteria"], test_layers, col_widths=[1.3, 1.2, 2.5, 1.5])

    add_heading_styled(doc, "12.2 Test Cases Summary", level=2)
    test_cases = [
        ["TC001", "User Registration", "FR1", "Register with valid email → Account created", "Pass"],
        ["TC002", "Duplicate Email Rejection", "FR1", "Register with existing email → Error message", "Pass"],
        ["TC003", "OTP Verification", "FR2", "Enter valid OTP within 5s → Account verified", "Pass"],
        ["TC004", "Account Lockout", "FR3", "5 failed logins → Account locked", "Pass"],
        ["TC005", "Password Reset", "FR5", "Reset with valid token (< 10 min) → Password updated", "Pass"],
        ["TC006", "Mood Entry Submission", "FR27", "Submit mood 1–10 → Entry recorded", "Pass"],
        ["TC007", "Duplicate Mood Prevention", "FR28", "Submit twice within 1 hour → Rejected", "Pass"],
        ["TC008", "Risk Alert Trigger", "FR30", "3 consecutive mood < 2 → Alert generated", "Pass"],
        ["TC009", "Appointment Booking", "FR9", "Book available slot → Appointment created", "Pass"],
        ["TC010", "Conflict Detection", "FR10", "Book occupied slot → Rejected", "Pass"],
        ["TC011", "Emergency Button", "FR38", "Press emergency → Crisis info < 1s", "Pass"],
        ["TC012", "Message Encryption", "FR23", "Send message → Stored encrypted (AES-256)", "Pass"],
        ["TC013", "Chatbot Response", "FR32", "Send query → Response < 3 seconds", "Pass"],
        ["TC014", "Chatbot Escalation", "FR33", "Low confidence → Escalated to clinician", "Pass"],
        ["TC015", "Consent Revocation", "FR54", "Revoke consent → Access immediately restricted", "Pass"],
        ["TC016", "Auto-Logout", "FR62", "15 min inactivity → Session terminated", "Pass"],
        ["TC017", "RBAC Enforcement", "FR65", "Patient access admin route → 403 Forbidden", "Pass"],
        ["TC018", "SQL Injection Prevention", "FR63", "Inject SQL in input → Sanitized/blocked", "Pass"],
        ["TC019", "File Upload Limit", "FR73", "Upload > 5MB file → Rejected", "Pass"],
        ["TC020", "Dashboard Load Time", "FR50", "Load dashboard → Render < 3 seconds", "Pass"],
    ]
    add_styled_table(doc, ["ID", "Test Case", "FR", "Procedure / Expected Result", "Status"], test_cases, col_widths=[0.6, 1.3, 0.5, 3.2, 0.5])

    add_heading_styled(doc, "12.3 Implementation Plan", level=2)
    impl_data = [
        ["Phase 1\nWeek 7", "Backend Foundation", "Express.js setup, PostgreSQL schema migration, JWT + OTP authentication, user CRUD APIs"],
        ["Phase 2\nWeek 8", "Core Feature APIs", "Mood tracking, appointment scheduling, messaging (AES-256), chatbot integration (OpenAI), emergency support APIs"],
        ["Phase 3\nWeek 9", "Frontend Development", "React component library, dashboard, mood tracker UI, messaging UI, appointment booking UI, admin panel"],
        ["Phase 4\nWeek 10", "Integration & Real-time", "Frontend-backend integration, Socket.io real-time chat, Google Calendar sync, notification engine, end-to-end testing"],
        ["Phase 5\nWeek 11", "Testing & Deployment", "Unit/integration/E2E testing, security audit, performance testing, cloud deployment (Render/AWS), CI/CD pipeline"],
        ["Phase 6\nWeek 12", "Demo & Handover", "Final demonstration, documentation completion, source code handover, deployment documentation"],
    ]
    add_styled_table(doc, ["Phase", "Focus", "Details"], impl_data, col_widths=[0.9, 1.5, 4.2])

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # 13. RESPONSES ON FEEDBACK
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "13. Responses on Feedback from Supervisor", level=1)
    add_body_text(doc, "This section documents the team's responses to feedback received from the academic supervisor, Mr. Syed Altaf, and the industry sponsor, Nabin Singh, during weekly project review sessions.")

    feedback_data = [
        ["Week 2", "Clarify the scope boundaries — what is in/out of scope for the prototype", "Added explicit scope and limitations section (Section 1.3) distinguishing prototype features from production-level features. Clearly stated that video consultation, mobile-native app, and real clinical diagnosis are out of scope."],
        ["Week 3", "Ensure functional requirements are measurable and testable", "Revised all 75 FRs to include specific measurable criteria (e.g., OTP within 5s, mood duplicate within 1hr, emergency response < 1s, API response < 500ms). Each FR now has a corresponding test case."],
        ["Week 4", "Architecture should clearly show separation of concerns", "Designed and documented three-tier architecture (Presentation, Application, Data) with clear component separation. Added detailed architecture diagram with technology labels for each layer."],
        ["Week 5", "Data Flow Diagrams should follow Gane & Sarson method consistently", "Redesigned all DFDs using Gane & Sarson notation. Created Level 0 Context Diagram, Level 1 DFD with 15 processes, and 15 Level 2 DFDs decomposing each process with proper data store references."],
        ["Week 5", "ERD should be database-ready with proper normalization", "Designed ERD with 15 normalized entities, UUIDs as primary keys, proper foreign key relationships, and cross-reference modulation. Created comprehensive data dictionary with constraints."],
        ["Week 6", "Include security considerations throughout the design", "Integrated DPIA (Section 11), implemented privacy controls table, added security-specific NFRs (NFRS01-06), and ensured all security FRs (FR23, FR54-55, FR62-66) are covered in DFDs and test cases."],
    ]
    add_styled_table(doc, ["Week", "Feedback", "Response / Action Taken"], feedback_data, col_widths=[0.7, 2.5, 3.5])

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "References", level=1)
    references = [
        "Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.",
        "Pressman, R. S. & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "IEEE (1998). IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications. IEEE.",
        "Gane, C. & Sarson, T. (1979). Structured Systems Analysis: Tools and Techniques. Prentice-Hall.",
        "OWASP Foundation (2021). OWASP Top 10 — 2021. https://owasp.org/Top10/",
        "React.js Documentation (2024). React — A JavaScript Library for Building User Interfaces. https://react.dev/",
        "Express.js Documentation (2024). Express — Fast, Unopinionated, Minimalist Web Framework. https://expressjs.com/",
        "PostgreSQL Documentation (2024). PostgreSQL: The World's Most Advanced Open Source Database. https://www.postgresql.org/docs/",
        "OpenAI API Documentation (2024). API Reference. https://platform.openai.com/docs/",
        "Google Calendar API Documentation (2024). Calendar API Overview. https://developers.google.com/calendar/",
        "Twilio Documentation (2024). Programmable Messaging. https://www.twilio.com/docs/",
        "SendGrid Documentation (2024). Email API. https://docs.sendgrid.com/",
        "Socket.IO Documentation (2024). Socket.IO. https://socket.io/docs/",
        "W3C (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
        "Information Commissioner's Office (ICO) (2022). Data Protection Impact Assessments. https://ico.org.uk/for-organisations/guide-to-data-protection/",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    # APPENDIX
    # ═══════════════════════════════════════════════
    add_heading_styled(doc, "Appendix", level=1)

    add_heading_styled(doc, "A. Table of Figures", level=2)
    figures = [
        ["Figure 1.1", "Organisation Chart — Mel23 Tech Solution"],
        ["Figure 1.2", "Work Breakdown Structure"],
        ["Figure 3.1", "Three-Tier System Architecture Diagram"],
        ["Figure 4.1", "Login / Registration Page Wireframe"],
        ["Figure 4.2", "User Dashboard Wireframe"],
        ["Figure 4.3", "Mood Tracker Page Wireframe"],
        ["Figure 4.4", "Secure Messaging Interface Wireframe"],
        ["Figure 4.5", "Appointments Page Wireframe"],
        ["Figure 4.6", "Admin Dashboard Panel Wireframe"],
        ["Figure 5.1", "Patient / User Use Case Diagram"],
        ["Figure 5.2", "Clinician / Doctor Use Case Diagram"],
        ["Figure 5.3", "Administrator Use Case Diagram"],
        ["Figure 6.1", "Level 0 Context Diagram (Gane & Sarson)"],
        ["Figure 7.1", "Level 1 System Data Flow Diagram"],
        ["Figure 7.2.1", "Level 2.1 — Authentication & Account Management DFD"],
        ["Figure 7.2.2", "Level 2.2 — Profile Management DFD"],
        ["Figure 7.2.3", "Level 2.3 — Appointment Management DFD"],
        ["Figure 7.2.4", "Level 2.4 — Clinical Records Management DFD"],
        ["Figure 7.2.5", "Level 2.5 — Messaging System DFD"],
        ["Figure 7.2.6", "Level 2.6 — Mood Tracking & Analysis DFD"],
        ["Figure 7.2.7", "Level 2.7 — AI Chatbot System DFD"],
        ["Figure 7.2.8", "Level 2.8 — Emergency Support DFD"],
        ["Figure 7.2.9", "Level 2.9 — Notification Engine DFD"],
        ["Figure 7.2.10", "Level 2.10 — Content & Resource Management DFD"],
        ["Figure 7.2.11", "Level 2.11 — Analytics & Reporting DFD"],
        ["Figure 7.2.12", "Level 2.12 — Administration & Access Control DFD"],
        ["Figure 7.2.13", "Level 2.13 — System Maintenance & Backup DFD"],
        ["Figure 7.2.14", "Level 2.14 — Privacy & Compliance DFD"],
        ["Figure 7.2.15", "Level 2.15 — UI & Localization DFD"],
        ["Figure 8.1", "SD1 — User Registration & Login Sequence Diagram"],
        ["Figure 8.2", "SD2 — Appointment Booking Flow Sequence Diagram"],
        ["Figure 8.3", "SD3 — Mood Entry & Risk Alert Sequence Diagram"],
        ["Figure 8.4", "SD4 — AI Chatbot Interaction Sequence Diagram"],
        ["Figure 8.5", "SD5 — Emergency Button Activation Sequence Diagram"],
        ["Figure 8.6", "SD6 — Secure Messaging Sequence Diagram"],
        ["Figure 9.1", "Entity Relationship Diagram"],
        ["Figure 10.1", "UML Class Diagram"],
    ]
    add_styled_table(doc, ["Figure #", "Description"], figures, col_widths=[1.2, 5.3])

    doc.add_paragraph()
    add_heading_styled(doc, "B. Table of Tables", level=2)
    tables_list = [
        ["Table 1.1", "Company Information"],
        ["Table 1.2", "Roles and Responsibilities"],
        ["Table 1.3", "Key Deliverables"],
        ["Table 1.4", "Project Budget Estimates"],
        ["Table 1.5", "WBS Phase Summary"],
        ["Table 1.6", "Gantt Chart / Project Timeline"],
        ["Table 2.1", "Functional Requirements (FR1–FR75)"],
        ["Table 2.2", "Non-Functional Requirements"],
        ["Table 3.1", "Technology Stack"],
        ["Table 7.1", "Level 1 Process Descriptions"],
        ["Table 7.2", "Data Store Definitions"],
        ["Table 9.1–9.11", "Data Dictionary Tables"],
        ["Table 11.1", "DPIA Risk Assessment"],
        ["Table 11.2", "Privacy Controls Summary"],
        ["Table 12.1", "Testing Strategy"],
        ["Table 12.2", "Test Cases Summary"],
        ["Table 12.3", "Implementation Plan"],
        ["Table 13.1", "Supervisor Feedback Responses"],
    ]
    add_styled_table(doc, ["Table #", "Description"], tables_list, col_widths=[1.2, 5.3])

    # ─── SAVE ───
    output_path = os.path.join(WORD_DIR, "Mental_Health_Support_Platform_Report.docx")
    # If file is locked (open in Word), save with suffix
    try:
        doc.save(output_path)
    except PermissionError:
        output_path = os.path.join(WORD_DIR, "Mental_Health_Support_Platform_Report_v2.docx")
        doc.save(output_path)
    print(f"\n✅ Report saved to: {output_path}")
    print(f"   File size: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == "__main__":
    generate_report()
