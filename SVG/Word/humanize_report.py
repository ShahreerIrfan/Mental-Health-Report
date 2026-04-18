"""
Humanize the Mental Health Support Platform Report docx.
Rewrites AI-sounding paragraphs with natural, student-written tone.
Preserves all formatting, images, tables, and styles.
"""
from docx import Document
from copy import deepcopy
import re, os

SRC = r"c:\Users\mdsha\Downloads\SVG\SVG\Word\Mental_Health_Support_Platform_Report_v2.docx"
DST = r"c:\Users\mdsha\Downloads\SVG\SVG\Word\Mental_Health_Support_Platform_Report_v2_Humanized.docx"

# Map: paragraph index -> new text (humanized)
# Only rewrite substantive prose paragraphs. Skip headings, figure captions,
# table labels, references, short items, and FR-number-only bullet lists.

REWRITES = {

    # === 1. Introduction ===
    66: (
        "Mental health plays a huge role in how people feel and function day to day, "
        "but getting the right support at the right time is not always easy. A lot of people "
        "struggle to keep track of how they are feeling or find trustworthy resources that "
        "can actually help. Things like stigma, not knowing where to look, and limited access "
        "to services make it even harder for people to reach out."
    ),
    67: (
        "This document is the Software Requirements Specification (SRS) and System Design "
        "report for our Mental Health Support Platform. It is a web-based app that lets users "
        "track their mood, browse helpful mental health resources, and use supportive tools — "
        "all within a secure and easy-to-use environment."
    ),

    # === 1.1 Purpose ===
    69: (
        "The main purpose of this project is to build a safe, easy-to-use web platform "
        "where people can look after their mental health. Users will be able to log their "
        "mood, write journal entries, and find useful mental health resources — all in one "
        "private and secure place."
    ),
    70: (
        "We want the platform to encourage self-awareness and help people catch warning "
        "signs early. Users can record how they feel each day, keep a personal journal, and "
        "browse organised materials like articles, wellness tips, and mental health support "
        "information. Over time, they will be able to spot patterns in their emotions."
    ),
    71: (
        "Security is a top priority. The system will use proper authentication and data "
        "protection so that users' personal information stays private and safe. On top of "
        "that, the platform will present data on a clean, simple dashboard to make it easy "
        "to use."
    ),
    72: (
        "We are following standard Software Development Lifecycle (SDLC) practices throughout "
        "this project — from gathering requirements and designing the system, through to "
        "building, testing, and deploying a working prototype."
    ),

    # === 1.2 Objectives ===
    74: (
        "The goal of this project is to design and build a working prototype that helps users "
        "keep track of their mental wellbeing and find helpful resources. Here are the main "
        "objectives we set out to achieve:"
    ),
    75: "Look into what users actually need when it comes to mental health tracking and support",
    77: "Build a full-stack web app using React for the front end, Node.js for the back end, and PostgreSQL for the database",
    79: "Set up secure login with JWT tokens, OTP verification, bcrypt password hashing, and AES-256 data encryption",
    80: "Give users access to a well-organised library of mental health resources",
    81: "Let patients and clinicians message each other securely in real time",
    82: "Add appointment booking that syncs with Google Calendar",
    83: "Get a working prototype deployed on a cloud hosting platform",

    # === 1.3 Scope & Limitations ===
    86: "Sign-up, login with JWT and OTP, and role-based access so different users see different things",
    87: "Daily mood logging on a 1–10 scale with notes and visual trend charts",
    88: "Private messaging between patients and clinicians, encrypted with AES-256",
    89: "Appointment booking with clash detection and Google Calendar sync",
    90: "An AI chatbot that uses NLP and can hand off to a real clinician if its confidence drops below 70%",
    91: "A one-tap emergency button that responds in under a second and shows crisis hotline details",
    92: "Clinical records that can only be viewed with proper consent",
    93: "A resource library where URLs are checked, content is grouped by topic, and views are counted",
    94: "An admin dashboard with analytics, audit logs, and compliance reports",
    95: "Privacy features like consent management and the option to delete your data",
    97: "This is a prototype — it is not meant to diagnose or treat any medical condition",
    98: "The AI chatbot gives general support only; it does not replace advice from a qualified professional",
    99: "Live video consultations are not included in this version",
    100: "All testing will use dummy data, not real patient records",
    101: "There is no native mobile app — the site will be responsive so it works on phones and tablets",

    # === 1.4 Assumptions ===
    103: "Users will be on a modern browser like Chrome, Firefox, Safari, or Edge",
    104: "A stable internet connection is needed for real-time features like chat",
    105: "Our team has access to the development tools and cloud services we need",
    106: "The third-party APIs we rely on (Google Calendar, OpenAI, Twilio, SendGrid) will stay available while we build",
    107: "We are using agile methodology and working in iterative sprints",
    108: "Every team member is available for the full 12-week project timeline",
    109: "Our academic supervisor and industry sponsor will give us feedback on time",
    110: "The PostgreSQL database will be hosted in the cloud with automatic backups turned on",

    # === 1.5 Company Profile ===
    113: (
        "Mel23 Tech Solution is our student development team, put together at Kent Institute "
        "of Technology Australia for the WIL (Work Integrated Learning) capstone project. We "
        "work under the guidance of our academic supervisor, Mr. Syed Altaf, and with "
        "sponsorship from Skillup Labs."
    ),
    115: (
        "Skillup Labs is the industry partner sponsoring this project. They help student teams "
        "build real-world tech products, especially in the HealthTech space."
    ),

    # === 1.6 Organisation Chart ===
    117: (
        "The chart below shows how our project team is structured. It includes the academic "
        "supervisor, industry sponsor, team leader, and each team member along with their role."
    ),

    # === 1.9 Budget ===
    123: (
        "Since this is a university capstone project, the budget mainly covers cloud hosting, "
        "third-party API costs, and development tools. The table below breaks down the "
        "estimated spending."
    ),

    # === 1.10 WBS ===
    125: (
        "The Work Breakdown Structure (WBS) splits the project into six main phases: "
        "Discovery, Requirements, Design, Development, Testing, and Deployment. Each phase "
        "has its own set of tasks and deliverables, as shown in the diagram below."
    ),

    # === 1.11 Gantt Chart ===
    129: (
        "The Gantt chart below maps out the project timeline, showing each phase and the "
        "key milestones across our 12-week development period."
    ),

    # === 2. System Requirements ===
    132: (
        "This section lays out what the platform needs to do (functional requirements) and "
        "how well it needs to do it (non-functional requirements). We gathered these through "
        "the project proposal, discussions with stakeholders, and looking at best practices "
        "in HealthTech app development."
    ),
    134: (
        "In total, we defined 75 functional requirements spread across modules like "
        "Authentication, Profile Management, Appointments, Clinical Records, Messaging, "
        "Mood Tracking, AI Chatbot, Emergency Support, Notifications, Content Management, "
        "Analytics, Administration, System Maintenance, Privacy, and UI/Localization."
    ),

    # === 3. System Architecture ===
    151: (
        "Our platform uses a three-tier architecture: a Presentation Tier (the front end), "
        "an Application Tier (the server), and a Data Tier (the database). This setup keeps "
        "things separated, makes the system easier to scale, and simpler to maintain."
    ),
    155: (
        "Presentation Tier (Client): This is the front end, built with React.js and Vite "
        "using TypeScript. It is a Single Page Application with component-based routing. "
        "We use React Context and React Query to manage state, Socket.io for real-time chat "
        "and notifications, Chart.js for mood analytics charts, and Axios with a JWT "
        "interceptor for secure API calls."
    ),
    156: (
        "Application Tier (Server): The back end runs on Express.js with Node.js. It has a "
        "REST API with middleware and error handling, an auth module that handles JWT tokens, "
        "OTP, bcrypt hashing, and role-based access. The business logic covers appointments, "
        "mood tracking, and clinical records. We also run a Socket.io server for real-time "
        "chat with AES-256 encryption, and connect to external services like OpenAI for the "
        "chatbot, Twilio for SMS, and SendGrid for email."
    ),
    157: (
        "Data Tier (Storage): PostgreSQL is our main database, storing users, appointments, "
        "clinical records, messages, and mood entries. We also use file storage for profile "
        "pictures, clinical documents, and backup files. Audit logs track all access, errors, "
        "and changes, while a caching layer handles JWT tokens, session data, and API "
        "response caching."
    ),
    158: (
        "Communication: The front end talks to the server over HTTPS and WSS (WebSocket "
        "Secure). The server communicates with the database through SQL queries via an ORM."
    ),

    # === 4. Wireframes ===
    163: (
        "We designed the wireframes below in Balsamiq Cloud. They show the main screens of "
        "the platform, including page layouts, navigation, and the key interactive elements."
    ),

    # === 5. Use Case Diagrams ===
    178: (
        "The use case diagrams below show how the three main actors — Patient/User, "
        "Clinician/Doctor, and Administrator — interact with the platform. Each use case "
        "links back to a specific functional requirement."
    ),

    # === 6. Context Diagram ===
    206: (
        "The Level 0 Context Diagram gives a bird's-eye view of the whole system and how it "
        "connects to outside entities. We used the Gane & Sarson style, where processes are "
        "shown as rounded rectangles and external entities are plain rectangles."
    ),
    210: "E1: Patient/User — The main user who signs up, logs mood, books appointments, and browses resources",
    211: "E2: Clinician — A health professional who handles appointments, writes session notes, and checks patient records",
    212: "E3: Administrator — The system admin who manages users, resources, analytics, and settings",
    213: "E4: SMS Gateway (Twilio) — An outside service we use for sending text messages and OTP codes",
    214: "E5: Google Calendar — Used to sync appointments with the user's calendar",
    215: "E6: Email Service (SendGrid) — Handles outgoing emails like booking confirmations and alerts",
    216: "E7: Crisis Hotline — An external service that provides emergency hotline information",

    # === 7. DFDs ===
    219: (
        "The Data Flow Diagrams (DFDs) break the system down level by level, showing how data "
        "moves between processes, outside entities, and data stores. Everything follows the "
        "Gane & Sarson notation."
    ),
    221: (
        "The Level 1 DFD takes the single system process and splits it into 15 sub-processes, "
        "each tied to specific functional requirements. It maps the data flows between "
        "external entities (E1–E7), processes (1.0–15.0), and 15 data stores (D1–D15)."
    ),
    228: (
        "The Level 2 DFDs drill deeper into each Level 1 process. They show the internal "
        "data flows, decision points, and which data stores are used within each process."
    ),
    229: (
        "Breaks down Process 1.0 into 7 steps: registering a user (FR1), verifying the OTP "
        "(FR2), locking accounts (FR3), logging login activity (FR4), resetting passwords "
        "(FR5), login/logout (FR8), and deleting and anonymising data (FR8d)."
    ),
    232: (
        "Breaks down Process 2.0 into 2 steps: updating profile details (FR6) and uploading "
        "a profile image with a file-size check under 2 MB (FR7)."
    ),
    235: (
        "Breaks down Process 3.0 into 9 steps covering things like scheduling (FR9), checking "
        "slot availability (FR10), rescheduling (FR11), marking missed appointments (FR12), "
        "pending approvals (FR14), auto-cancel (FR15), Google Calendar sync (FR67), follow-up "
        "notifications (FR68), and rebooking from alerts (FR69)."
    ),
    238: (
        "Breaks down Process 4.0 into 10 steps: writing session notes (FR16), checking consent "
        "(FR17), viewing patient history (FR18), managing record retention (FR19), reassigning "
        "patients (FR20), handling medical record requests (FR21), logging access (FR22), "
        "uploading clinical documents (FR51), validating file formats (FR52), and downloading "
        "records (FR53)."
    ),
    241: (
        "Breaks down Process 5.0 into 4 steps: encrypting messages with AES-256 (FR23), "
        "delivering them within 2 seconds (FR24), storing chat transcripts (FR25), and soft "
        "deleting messages (FR26)."
    ),
    244: (
        "Breaks down Process 6.0 into 5 steps: submitting a mood entry (FR27), preventing "
        "duplicates within a 1-hour window (FR28), generating weekly mood reports (FR29), "
        "triggering risk alerts when scores are low several days in a row (FR30), and "
        "recommending resources based on mood patterns (FR31)."
    ),
    247: (
        "Breaks down Process 7.0 into 5 steps: processing a query within 3 seconds (FR32), "
        "escalating to a human when confidence is below 70% (FR33), saving transcripts "
        "(FR34), letting admins review logs (FR35), and flagging users who have been inactive "
        "for 7 days (FR36)."
    ),
    250: (
        "Breaks down Process 8.0 into 3 steps: activating the emergency button in under 1 "
        "second (FR38), showing crisis hotline info (FR39), and logging the activation (FR40)."
    ),
    253: (
        "Breaks down Process 9.0 into 6 steps: sending a booking confirmation (FR13), "
        "dispatching wellness checks (FR37), integrating SMS and email (FR57), retrying "
        "failed sends up to 3 times (FR58), queuing messages during downtime (FR59), and "
        "following up on missed appointments (FR68)."
    ),
    256: (
        "Breaks down Process 10.0 into 3 steps: adding resources through the admin panel "
        "(FR41), checking URLs before they go live (FR42), and tracking how many times each "
        "resource is viewed (FR43)."
    ),
    259: (
        "Breaks down Process 11.0 into 6 steps: generating monthly reports (FR44), tracking "
        "session times (FR61), evaluating clinician workloads (FR70), producing compliance "
        "reports (FR71), logging data changes (FR72), and exporting to CSV (FR74)."
    ),
    262: (
        "Breaks down Process 12.0 into 6 steps: deactivating accounts (FR60), auto-logging "
        "out after 15 minutes of inactivity (FR62), validating inputs to stop SQL injection "
        "(FR63), creating audit trails (FR64), assigning roles via RBAC (FR65), and blocking "
        "unauthorised API requests (FR66)."
    ),
    265: (
        "Breaks down Process 13.0 into 6 steps: running a daily backup at midnight (FR45), "
        "restoring backups (FR46), logging system errors (FR47), alerting admins about "
        "critical issues (FR48), capping file uploads at 5 MB (FR73), and scheduling monthly "
        "maintenance (FR75)."
    ),
    268: (
        "Breaks down Process 14.0 into 3 steps: letting users revoke consent instantly "
        "(FR54), cutting off access as soon as consent is withdrawn (FR55), and logging every "
        "consent change (FR56)."
    ),
    271: (
        "Breaks down Process 15.0 into 2 steps: choosing an interface language (FR49) and "
        "making sure the dashboard loads in under 3 seconds (FR50)."
    ),

    # === 8. Sequence Diagrams ===
    276: (
        "The sequence diagrams show how different parts of the system talk to each other "
        "during key actions. They map out the order of messages between actors, the front "
        "end, the API server, the database, and any outside services."
    ),
    277: (
        "This diagram walks through the full sign-up process: the user submits their details, "
        "a record is created in the database, an OTP is sent by email, the user enters the "
        "code, their account gets verified, a JWT token is issued, and the dashboard loads. "
        "The participants are the Patient, Frontend, API Server, and Database."
    ),
    280: (
        "This one covers appointment booking from start to finish: the user picks a clinician "
        "and date, the system checks if the slot is free, confirms the booking, saves it to "
        "the database, syncs with Google Calendar, and sends out an email notification. "
        "Participants include the Patient, Frontend, API Server, Database, and Google Calendar."
    ),
    283: (
        "Here we show the mood submission flow. The system checks for duplicate entries "
        "within the last 24 hours, saves the mood score, looks for consecutive low scores "
        "(below 2 for 3 days straight), fires a risk alert if needed, logs a notification, "
        "and updates the dashboard."
    ),
    286: (
        "This diagram covers a chatbot conversation: the user sends a message, the API "
        "server processes it, sends an NLP query to the AI engine, gets a confidence score "
        "back, delivers the response in under 3 seconds, saves the transcript, and — if "
        "confidence is below 70% — automatically passes the conversation to a real clinician."
    ),
    289: (
        "This shows what happens when someone hits the emergency button: the system calls "
        "the emergency API, creates a log entry, fetches crisis hotline data from an outside "
        "service, displays the information in under 1 second, and records everything in the "
        "audit trail. Participants are the Patient, Frontend, API Server, Database, and "
        "Crisis Service."
    ),
    292: (
        "This diagram follows a secure message from start to finish: the user writes the "
        "message, the server encrypts it with AES-256, stores the encrypted version, pushes "
        "it to the recipient in real time via WebSocket (under 2 seconds), confirms delivery, "
        "sends a notification, and supports soft deletion."
    ),

    # === 9. ERD ===
    297: (
        "The Entity Relationship Diagram (ERD) shows the logical data model behind our "
        "database. It lays out each table, its primary keys (PK) and foreign keys (FK), and "
        "how the tables relate to each other with clear cardinality."
    ),
    301: "Below are the schema definitions for every entity in the database:",

    # === 10. Class Diagram ===
    326: (
        "The UML Class Diagram maps out the object-oriented structure of the platform. Each "
        "class lists its attributes (properties) and methods (what it can do), with lines "
        "showing how the classes are linked through associations, dependencies, and "
        "multiplicities."
    ),
    330: "User — Handles sign-up, login, OTP checks, password resets, and account deletion",
    331: "Profile — Manages the user's display name, profile picture, and language setting",
    332: "Appointment — Covers scheduling, clash detection, slot checking, rescheduling, approvals, auto-cancel, and missed marking",
    333: "ClinicalRecord — Deals with medical records, consent checks, session notes, document uploads and downloads, and viewing patient history",
    334: "ChatMessage — Sends and receives encrypted messages (AES-256) with delivery tracking and soft delete",
    335: "MoodEntry — Tracks daily moods, stops duplicate entries, builds weekly reports, raises risk alerts, and suggests resources",
    336: "Chatbot — Runs AI conversations with NLP, can escalate to a human, stores transcripts, and spots inactive users",
    337: "Emergency — Activates the emergency feature, pulls up hotline data, and logs each activation",
    338: "Notification — Sends alerts through email, SMS, and in-app channels with retry logic and downtime queuing",
    339: "Resource — Manages the mental health resource library, checks URLs, and tracks view counts",
    340: "Analytics — Produces monthly reports, tracks clinician workloads, generates compliance reports, and exports CSV files",
    341: "Admin — Gives admins control over user management, role assignments, account deactivation, audit logs, and backups",

    # === 11. DPIA ===
    344: (
        "The Data Protection Impact Assessment (DPIA) looks at the privacy risks that come "
        "with handling personal and sensitive health information on our platform, and explains "
        "the steps we take to reduce those risks."
    ),
    347: "Account Data — Things like email addresses, hashed passwords, login times, and IP addresses",
    348: "Profile Data — Display names, profile pictures, and language choices",
    349: "Health Data — Mood scores, mood notes, clinical session notes, and medical records",
    350: "Communication Data — Encrypted chat messages and chatbot conversation logs",
    351: "Behavioural Data — How long sessions last, usage patterns, and resource view counts",
    352: "Emergency Data — Logs from emergency activations and crisis interactions",
    353: "Consent Data — What data-sharing choices the user has made and any changes over time",

    # === 12. Testing ===
    359: (
        "Our testing strategy uses multiple layers to make sure every part of the system "
        "works properly:"
    ),

    # === 13. Supervisor Feedback ===
    364: (
        "In this section, we document how the team responded to feedback from our academic "
        "supervisor, Mr. Syed Altaf, and our industry sponsor, Nabin Singh, during the "
        "weekly project review meetings."
    ),
}


def rewrite_paragraph(para, new_text):
    """
    Replace paragraph text while keeping the formatting of the first run.
    This preserves font, size, bold, italic, color, etc.
    """
    if not para.runs:
        para.text = new_text
        return

    # Keep first run's formatting, put all new text in it
    first_run = para.runs[0]
    first_run.text = new_text

    # Remove all other runs
    for run in para.runs[1:]:
        run.text = ""


def main():
    print(f"Opening: {SRC}")
    doc = Document(SRC)

    count = 0
    for idx, new_text in REWRITES.items():
        if idx < len(doc.paragraphs):
            old_text = doc.paragraphs[idx].text.strip()
            if old_text:
                rewrite_paragraph(doc.paragraphs[idx], new_text)
                count += 1
                preview = new_text[:80] + "..." if len(new_text) > 80 else new_text
                print(f"  [Para {idx}] Rewritten: {preview}")

    doc.save(DST)
    print(f"\nDone! Rewrote {count} paragraphs.")
    print(f"Saved to: {DST}")


if __name__ == "__main__":
    main()
