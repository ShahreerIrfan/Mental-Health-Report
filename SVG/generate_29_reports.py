#!/usr/bin/env python3
"""
Generate 29 separate DOCX files from the Mental Health Support Platform report.
Uses pre-cached high-res PNGs from Word/_png_cache/.
"""

import os, io, glob
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE_DIR = r"c:\Users\mdsha\Downloads\SVG\SVG"
PNG_CACHE = os.path.join(BASE_DIR, "Word", "_png_cache")
OUT_DIR = os.path.join(BASE_DIR, "Report_Sections")
os.makedirs(OUT_DIR, exist_ok=True)

# ════════════════════════════════════════════════════════
# STYLING HELPERS (improved design)
# ════════════════════════════════════════════════════════

BLUE_DARK = RGBColor(0x0D, 0x47, 0x71)
BLUE_MED = RGBColor(0x1A, 0x6F, 0xA3)
BLUE_LIGHT = RGBColor(0x29, 0x80, 0xB9)
GREY_DARK = RGBColor(0x2C, 0x3E, 0x50)
GREY_MID = RGBColor(0x4A, 0x4A, 0x4A)
GREY_LIGHT = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def new_doc():
    """Create a fresh document with improved styling."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = GREY_DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(16)
    return doc


def add_section_header(doc, section_num, section_title, subtitle=None):
    """Add a beautiful section title page."""
    for _ in range(5):
        doc.add_paragraph()

    # Decorative line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 55)
    r.font.color.rgb = BLUE_LIGHT
    r.font.size = Pt(11)

    # Section number
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(section_num)
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE_LIGHT
    r.bold = True

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(section_title)
    r.font.size = Pt(26)
    r.font.color.rgb = BLUE_DARK
    r.bold = True

    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        r.font.size = Pt(12)
        r.font.color.rgb = GREY_LIGHT
        r.italic = True

    # Line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 55)
    r.font.color.rgb = BLUE_LIGHT
    r.font.size = Pt(11)

    for _ in range(3):
        doc.add_paragraph()

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mental Health Support Platform")
    r.font.size = Pt(13)
    r.font.color.rgb = GREY_MID
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mel23 Tech Solution  |  Kent Institute of Technology Australia")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_LIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("April 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY_LIGHT

    doc.add_page_break()


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE_DARK
    return h


def body(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    r.font.color.rgb = GREY_DARK
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = Pt(18)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    for r in p.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = GREY_DARK
    return p


def table(doc, headers, rows, col_widths=None, hdr_color="0D4771"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_shading(cell, hdr_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Calibri'
            r.font.color.rgb = GREY_DARK
            if ri % 2 == 1:
                set_shading(cell, "EAF2F8")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def img(doc, svg_name, caption=None, width=6.0):
    """Insert cached PNG image."""
    png = os.path.join(PNG_CACHE, svg_name.replace('.svg', '.png'))
    if os.path.exists(png):
        doc.add_picture(png, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Diagram: {svg_name}]")
        r.italic = True
        r.font.color.rgb = GREY_LIGHT
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY_MID
        p.paragraph_format.space_after = Pt(10)


def save(doc, filename):
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    sz = os.path.getsize(path) / 1024
    print(f"  ✓ {filename} ({sz:.0f} KB)")


def add_footer_text(doc, text):
    """Add a small footer-style note at the bottom."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_LIGHT
    r.italic = True


# ════════════════════════════════════════════════════════
# 29 DOCUMENT GENERATORS
# ════════════════════════════════════════════════════════

def gen_00_cover_toc():
    doc = new_doc()
    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Mental Health Support Platform")
    r.font.size = Pt(32)
    r.font.color.rgb = BLUE_DARK
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WIL Capstone Project Report")
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE_MED

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 55)
    r.font.color.rgb = BLUE_LIGHT

    cover_info = [
        ("Team", "Mel23 Tech Solution"),
        ("Institution", "Kent Institute of Technology Australia"),
        ("Academic Supervisor", "Mr. Syed Altaf"),
        ("Industry Sponsor", "Skillup Labs (Nabin Singh)"),
        ("Submission Date", "April 2026"),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  ")
        r.font.size = Pt(12)
        r.font.color.rgb = GREY_MID
        r = p.add_run(value)
        r.font.size = Pt(12)
        r.font.color.rgb = BLUE_DARK
        r.bold = True

    doc.add_page_break()

    heading(doc, "Table of Contents")
    toc_items = [
        ("01", "1.1  Purpose of the Project"),
        ("02", "1.2  Objective of the Project"),
        ("03", "1.3  Project Scope and Limitations"),
        ("04", "1.4  Assumptions"),
        ("05", "1.5  Company Profile"),
        ("06", "1.6  Organisation Chart"),
        ("07", "1.7  Roles and Responsibilities"),
        ("08", "1.8  Key Deliverables"),
        ("09", "1.9  Project Budget"),
        ("10", "1.10  Work Breakdown Structure"),
        ("11", "1.11  Gantt Chart / Project Timeline"),
        ("12", "2.1  Functional Requirements (FR1–FR75)"),
        ("13", "2.2  Non-Functional Requirements"),
        ("14", "3.1  Software & Hardware Architecture"),
        ("15", "3.2  Tech Stack"),
        ("16", "4.1  User Interface Design / Wireframes"),
        ("17", "5.  Use Case Diagrams"),
        ("18", "6.  Context Diagram (Level 0)"),
        ("19", "7.1  DFD Level 1"),
        ("20", "7.2  DFD Level 2"),
        ("21", "8.  Sequence Diagrams"),
        ("22", "9.  ERD & Data Dictionary"),
        ("23", "10.  UML Class Diagram"),
        ("24", "11.  Data Protection Impact Assessment"),
        ("25", "12.  Test and Implementation Plan"),
        ("26", "13.  Responses on Supervisor Feedback"),
        ("27", "References"),
        ("28", "Appendix"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        r = p.add_run(f"  File {num}  ")
        r.font.size = Pt(10)
        r.font.color.rgb = GREY_LIGHT
        r.font.name = 'Calibri'
        r = p.add_run(f"  {title}")
        r.font.size = Pt(11)
        r.font.color.rgb = GREY_DARK
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(3)

    add_footer_text(doc, "Mental Health Support Platform — Cover Page & Table of Contents")
    save(doc, "00_Cover_and_Table_of_Contents.docx")


def gen_01_purpose():
    doc = new_doc()
    add_section_header(doc, "Section 1.1", "Purpose of the Project", "Why We Are Building This Platform")
    heading(doc, "1.1 Purpose of the Project")
    body(doc, "Mental health is something that affects each and every one of us, yet so many people are not able to get the help they need at the right time. There are many reasons for this — stigma around mental health, not knowing where to look for help, and honestly, most of the existing solutions are either too expensive or too complicated for a common person to use.")
    body(doc, "Keeping all of this in mind, the purpose of our project is to design and build a safe, easy-to-use, web-based Mental Health Support Platform. The idea is simple — we want to give users a private space where they can keep track of how they are feeling on a daily basis, write down their thoughts through journaling, and access properly curated mental health resources whenever they need them.")
    body(doc, "We are not trying to replace professional therapists or doctors here — that is not the goal at all. What we want is to help people become more self-aware about their emotional patterns. When someone tracks their mood every day, they start noticing things — like what triggers a bad day, or what helps them feel better. This kind of self-awareness is actually the first step towards getting better.")
    body(doc, "The platform will also have features like an AI-powered chatbot that can provide immediate support when a user is feeling low, a secure messaging system so that patients can communicate with clinicians privately, and an emergency button that shows crisis hotline numbers instantly when someone is in a really bad state.")
    body(doc, "From a technical standpoint, we are building this with proper security measures in place. All the sensitive data — mood entries, chat messages, clinical records — everything will be encrypted. We are using AES-256 encryption for messages, bcrypt for password hashing, and JWT tokens for authentication. Privacy is not just a feature for us, it is a core principle of the entire system.")
    body(doc, "At the end of the day, this platform is our attempt to make mental health support more accessible, more affordable, and more convenient for everyday people. Even if it helps a handful of users to feel better about themselves, we would consider that a success.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.1 Purpose of the Project")
    save(doc, "01_Purpose_of_the_Project.docx")


def gen_02_objective():
    doc = new_doc()
    add_section_header(doc, "Section 1.2", "Objective of the Project", "What We Aim to Achieve")
    heading(doc, "1.2 Objective of the Project")
    body(doc, "The main objective of this project is to design and develop a working prototype of a web-based platform that genuinely helps users to take care of their mental wellbeing. We did not want to just build something for the sake of completing an assignment — we wanted to create something that actually makes a difference, even if it is on a small scale.")
    body(doc, "Given below are the specific objectives that we have set for ourselves:")
    objectives = [
        "To properly analyse and understand what users actually need when it comes to mental health tracking and support — not just assumptions, but real needs based on research",
        "To define clear and measurable system requirements that cover everything from basic user registration to advanced AI chatbot interactions",
        "To design and implement a full-stack web application using React.js for the frontend, Node.js with Express.js for the backend, and PostgreSQL for our database",
        "To develop core features including daily mood tracking with a 1–10 scale, personal journaling, and trend visualisation using charts so users can actually see how their mood changes over time",
        "To build a secure authentication system using JWT tokens, OTP verification, and bcrypt password hashing — because when you are dealing with mental health data, security cannot be an afterthought",
        "To integrate an AI chatbot powered by OpenAI that can respond to user queries within 3 seconds and automatically escalate to a human clinician if the confidence level drops below 70%",
        "To implement real-time secure messaging between patients and clinicians using Socket.io with AES-256 encryption — no one should be able to read these conversations except the sender and receiver",
        "To enable appointment scheduling with proper slot conflict detection and synchronisation with Google Calendar so that nothing gets double-booked",
        "To provide a well-organised resource library with curated mental health articles, tips, and support information that users can browse through at their own pace",
        "To deploy a working prototype on a cloud hosting environment and make sure it performs reliably under expected user loads",
    ]
    for o in objectives:
        bullet(doc, o)
    body(doc, "These objectives are not just theoretical goals written for documentation purposes. Each one of them has corresponding functional requirements (FR1 through FR75) and test cases that we will be using to verify whether we have actually achieved what we set out to do.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.2 Objective of the Project")
    save(doc, "02_Objective_of_the_Project.docx")


def gen_03_scope():
    doc = new_doc()
    add_section_header(doc, "Section 1.3", "Project Scope and Limitations", "What Is Included and What Is Not")
    heading(doc, "1.3 Project Scope and Limitations")
    body(doc, "Before jumping into development, it is really important that we define clearly what this project will cover and what it will not. This helps us manage expectations properly — both for our team and for anyone evaluating this project.")
    heading(doc, "What Is Within Scope", level=2)
    body(doc, "The following features and functionalities are included in this project:")
    scope_items = [
        "User registration and authentication with email/SIM validation, OTP verification, and role-based access control (RBAC) supporting three roles — Patient, Clinician, and Administrator",
        "Daily mood tracking where users can rate their mood on a 1–10 scale, add personal notes, and view trend visualisations through interactive charts",
        "Secure messaging system with end-to-end AES-256 encryption between patients and clinicians, with message delivery guaranteed within 2 seconds",
        "Appointment scheduling system with proper slot conflict detection, automatic reminders, and synchronisation with Google Calendar",
        "AI-powered chatbot that uses OpenAI's NLP engine to provide immediate responses, with automatic escalation to a human clinician when confidence drops below 70%",
        "Emergency support button that activates within less than 1 second and immediately displays crisis hotline numbers and support contact information",
        "Clinical records management where clinicians can write session notes, upload documents, and view patient history — all controlled through consent-based access",
        "Resource library managed by administrators, with URL validation before publishing and view count tracking for analytics",
        "Administrative dashboard with user management, analytics reporting, audit trails, and regulatory compliance reports",
        "Privacy controls including granular consent management, right to data deletion, and data anonymisation for analytics purposes",
    ]
    for s in scope_items:
        bullet(doc, s)

    heading(doc, "What Is Outside Scope (Limitations)", level=2)
    body(doc, "We want to be upfront about what this project does not include:")
    limitations = [
        "This platform is strictly a prototype — it is not meant for actual clinical diagnosis or treatment of any mental health condition",
        "The AI chatbot provides supplementary support only. Its responses should never be treated as a substitute for professional medical advice",
        "Real-time video consultation between patients and clinicians is not part of the current scope. We considered it, but it would have added too much complexity for the given timeline",
        "The system will be tested using simulated data, not real patient records. We are not in a position to handle actual medical data at this stage",
        "A native mobile application (iOS/Android) is not being developed. However, the web application will be fully responsive and work well on mobile browsers",
    ]
    for l in limitations:
        bullet(doc, l)
    add_footer_text(doc, "Mental Health Support Platform — Section 1.3 Project Scope and Limitations")
    save(doc, "03_Scope_and_Limitations.docx")


def gen_04_assumptions():
    doc = new_doc()
    add_section_header(doc, "Section 1.4", "Assumptions", "What We Are Taking for Granted")
    heading(doc, "1.4 Assumptions")
    body(doc, "Every project operates on certain assumptions, and it is better to state them clearly rather than leaving them implied. The following assumptions have been made while planning and developing this platform:")
    assumptions = [
        "Users will be accessing the platform through a modern web browser such as Google Chrome, Mozilla Firefox, Safari, or Microsoft Edge. We are not providing support for outdated browsers like Internet Explorer",
        "Users will have a reasonably stable internet connection, especially for real-time features like chat messaging and live notifications. If the connection drops, queued messages will be delivered once connectivity is restored",
        "All team members will have access to the required development tools — VS Code, Node.js, PostgreSQL, Git — and cloud services like Render or AWS for deployment",
        "Third-party APIs that we are integrating with — Google Calendar, OpenAI, Twilio for SMS, and SendGrid for emails — will remain available and functional throughout the development period",
        "The team will follow agile methodology with iterative sprints, weekly stand-ups, and regular code reviews through GitHub pull requests",
        "All six team members will be available and actively contributing throughout the 12-week project duration",
        "Our academic supervisor, Mr. Syed Altaf, and our industry sponsor, Mr. Nabin Singh from Skillup Labs, will provide timely feedback during weekly review sessions",
        "The PostgreSQL database will be hosted on a cloud provider with automated daily backups and a 7-day retention policy, so we do not lose any data",
    ]
    for a in assumptions:
        bullet(doc, a)
    body(doc, "If any of these assumptions turn out to be incorrect during the project, we will reassess and adjust our plans accordingly. That is the beauty of working with agile — we can adapt.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.4 Assumptions")
    save(doc, "04_Assumptions.docx")


def gen_05_company_profile():
    doc = new_doc()
    add_section_header(doc, "Section 1.5", "Company Profile", "About Our Team and Industry Sponsor")
    heading(doc, "1.5 Company Profile")
    heading(doc, "Mel23 Tech Solution", level=2)
    body(doc, "Mel23 Tech Solution is our student-led development team that was formed specifically for this WIL (Work Integrated Learning) capstone project at Kent Institute of Technology Australia. The name 'Mel23' comes from the fact that we are based in Melbourne and our batch started in 2023.")
    body(doc, "Our team consists of six members, each bringing a different set of skills to the table. We have a team leader who handles project coordination and stakeholder communication, a full-stack tech lead who takes care of the overall architecture and code quality, a dedicated backend developer, a frontend developer, a QA engineer who makes sure nothing breaks, and a UI/UX designer who ensures the platform looks good and is easy to use.")
    body(doc, "We operate under the academic guidance of Mr. Syed Altaf, who has been our supervisor throughout this journey. His feedback during weekly review sessions has been instrumental in shaping the direction of this project.")

    heading(doc, "Industry Sponsor — Skillup Labs", level=2)
    body(doc, "Skillup Labs is the industry project provider that sponsors our Mental Health Support Platform project. They are an organisation that actively supports student teams in building real-world technology solutions, particularly within the HealthTech domain.")
    body(doc, "Mr. Nabin Singh from Skillup Labs has been our primary point of contact. He provided us with the initial project brief and has been available for consultations whenever we needed industry-level guidance on requirements and feasibility.")

    table(doc, ["Attribute", "Details"], [
        ["Industry Provider", "Skillup Labs"],
        ["Program Contact", "Nabin Singh"],
        ["Contact Email", "wil@skilluplabs.com.au"],
        ["Academic Supervisor", "Mr. Syed Altaf"],
        ["Institution", "Kent Institute of Technology Australia"],
        ["Project Duration", "12 Weeks (Trimester-based)"],
        ["Team Name", "Mel23 Tech Solution"],
        ["Team Size", "6 Members"],
    ], col_widths=[2.5, 4.0])
    add_footer_text(doc, "Mental Health Support Platform — Section 1.5 Company Profile")
    save(doc, "05_Company_Profile.docx")


def gen_06_org_chart():
    doc = new_doc()
    add_section_header(doc, "Section 1.6", "Organisation Chart", "Team Structure and Hierarchy")
    heading(doc, "1.6 Organisation Chart")
    body(doc, "The organisation chart below shows how our team is structured. It is a straightforward hierarchy — the academic supervisor and industry sponsor sit at the top providing guidance, followed by our team leader who coordinates everything, and then the five team members who handle the actual development and design work.")
    body(doc, "We deliberately kept the structure flat because we felt that in a team of six people, having too many layers of hierarchy would just slow things down. Everyone reports to the team leader, and the team leader reports to both the academic supervisor and the industry sponsor.")
    img(doc, "1.6 Organisation chart_diagram_1.svg", "Figure 1.1: Organisation Chart — Mel23 Tech Solution", 6.2)
    body(doc, "This structure has worked quite well for us so far. Communication flows freely, and when there is a blocker, it gets resolved quickly because there is no unnecessary chain of command to go through.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.6 Organisation Chart")
    save(doc, "06_Organisation_Chart.docx")


def gen_07_roles():
    doc = new_doc()
    add_section_header(doc, "Section 1.7", "Roles and Responsibilities", "Who Does What in Our Team")
    heading(doc, "1.7 Roles and Responsibilities")
    body(doc, "Each team member has been assigned a specific role based on their strengths and interests. We made sure that the workload is distributed evenly and that everyone has a clear understanding of what they are responsible for. Given below is the detailed breakdown:")
    table(doc, ["Name / Student ID", "Role", "Key Responsibilities"], [
        ["Mr. Syed Altaf", "Academic Supervisor", "Provides academic guidance and direction, reviews all major deliverables, evaluates team progress during weekly sessions, and ensures the project meets the learning outcomes defined by the institution"],
        ["Nabin Singh\nSkillup Labs", "Industry Sponsor", "Provides real-world industry context, defines project requirements from a business perspective, reviews the final product against industry expectations, and gives feedback on feasibility"],
        ["Sanjana Tanwar\nK240225", "Team Leader /\nProject Manager", "Handles overall project coordination, sprint planning, task assignment using Jira or Trello, stakeholder communication, progress tracking, and makes sure the team meets all deadlines"],
        ["Jubair Zaman Dipto\nK240568", "Full-Stack Developer /\nTech Lead", "Responsible for the overall system architecture, API design, full-stack integration between frontend and backend, code reviews, and making key technical decisions for the team"],
        ["S M Nalid Maola\nK231528", "Backend Developer", "Develops all REST API endpoints using Express.js, designs and maintains the PostgreSQL database schema, implements business logic, and handles authentication and security modules"],
        ["Nabil Ashrafi\nK231720", "Frontend Developer", "Builds all React.js components and pages, implements responsive UI using Tailwind CSS, manages client-side state with React Context and React Query, and integrates with backend APIs"],
        ["Md Al Amin Sikder\nK240094", "QA Engineer /\nTester", "Creates and executes test plans, writes unit tests using Jest, integration tests with Supertest, end-to-end tests using Cypress, manages the CI/CD pipeline with GitHub Actions"],
        ["Tranpreet Kaur Kumar\nK241957", "UI/UX Designer", "Designs wireframes using Balsamiq Cloud, conducts basic UX research, ensures WCAG 2.1 Level AA accessibility compliance, and creates design documentation for the team to follow"],
    ], col_widths=[1.6, 1.3, 3.8])
    body(doc, "We have weekly stand-up meetings every Monday where each member gives a quick update on what they have done, what they plan to do next, and if there are any blockers. This has been really helpful in keeping everyone on the same page.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.7 Roles and Responsibilities")
    save(doc, "07_Roles_and_Responsibilities.docx")


def gen_08_deliverables():
    doc = new_doc()
    add_section_header(doc, "Section 1.8", "Key Deliverables", "What We Will Hand Over at the End")
    heading(doc, "1.8 Key Deliverables")
    body(doc, "Throughout this 12-week project, we are committed to producing a set of deliverables that demonstrate the complete software development lifecycle — from initial analysis all the way to deployment and final presentation. Each deliverable has a clear deadline and a responsible person.")
    table(doc, ["ID", "Deliverable", "Due", "Responsible"], [
        ["D1", "Problem Analysis Report — understanding the problem space, user personas, market research", "Week 2", "Sanjana (TL)"],
        ["D2", "Software Requirements Specification (SRS) — all 75 functional and non-functional requirements", "Week 3", "Nalid + Dipto"],
        ["D3", "System Architecture Documentation — three-tier architecture, API specifications, tech stack decisions", "Week 4", "Dipto (Tech Lead)"],
        ["D4", "Database Schema Design (ERD) — normalised entity relationship diagram with 15 entities", "Week 5", "Nalid"],
        ["D5", "UI/UX Wireframes — 6 key screens designed in Balsamiq Cloud", "Week 5", "Tranpreet"],
        ["D6", "Functional Web Application Prototype — working application with core features", "Week 9", "Full Team"],
        ["D7", "Source Code Repository — well-documented GitHub repository with proper branching strategy", "Ongoing", "Dipto (Tech Lead)"],
        ["D8", "Testing Documentation — unit, integration, E2E, and security test results", "Week 11", "Al Amin"],
        ["D9", "Deployment Documentation — cloud deployment guide, CI/CD pipeline configuration", "Week 11", "Dipto + Al Amin"],
        ["D10", "Final Presentation and Live Demonstration — showcasing the complete system to evaluators", "Week 12", "Sanjana + Team"],
    ], col_widths=[0.5, 3.3, 0.8, 1.5])
    body(doc, "We are tracking all deliverables through our project management tool and making sure nothing slips through the cracks. Every deliverable goes through at least one round of peer review before it is considered final.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.8 Key Deliverables")
    save(doc, "08_Key_Deliverables.docx")


def gen_09_budget():
    doc = new_doc()
    add_section_header(doc, "Section 1.9", "Project Budget", "Cost Estimates and Resource Allocation")
    heading(doc, "1.9 Project Budget")
    body(doc, "Since this is an academic capstone project, we do not have a massive budget to work with. However, we have carefully estimated the costs involved, and thankfully, most of the tools and services we need either have free tiers or offer student discounts. The detailed budget breakdown is given below:")
    table(doc, ["Item", "Estimated Cost", "Notes"], [
        ["Cloud Hosting (Render / AWS)", "$0 – $50/month", "We are planning to use Render's free tier for initial deployment. If we need more resources, AWS student credits will cover it"],
        ["PostgreSQL Database Hosting", "$0 – $20/month", "Render or Supabase free tier should be sufficient for our prototype's data volume"],
        ["Domain Name (optional)", "$10 – $15/year", "We might register a .com domain for the final demo, but this is optional"],
        ["OpenAI API (AI Chatbot)", "$10 – $30/month", "GPT-3.5 Turbo is quite affordable. We will set usage limits to control costs during development"],
        ["Twilio SMS API", "$0 – $15/month", "Trial credits are available which should last through the development period"],
        ["SendGrid Email API", "$0", "Free tier gives us 100 emails per day, which is more than enough for our prototype"],
        ["Google Calendar API", "$0", "Completely free for our usage level"],
        ["Balsamiq Cloud (Wireframes)", "$0", "Educational license provided by the institution"],
        ["GitHub (Version Control)", "$0", "Free for students through GitHub Education"],
        ["Development Tools", "$0", "VS Code, Postman, pgAdmin — all free and open-source"],
        ["Total Monthly Estimate", "$20 – $130/month", "Most costs are variable and usage-based"],
    ], col_widths=[2.2, 1.3, 3.2])
    body(doc, "We are being quite conservative with the budget. The reality is that for a prototype project like ours, free tiers and student credits cover almost everything. The only significant expense would be OpenAI API usage, and even that can be kept under control by using GPT-3.5 Turbo instead of GPT-4 for most interactions.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.9 Project Budget")
    save(doc, "09_Project_Budget.docx")


def gen_10_wbs():
    doc = new_doc()
    add_section_header(doc, "Section 1.10", "Work Breakdown Structure", "How We Have Divided the Work")
    heading(doc, "1.10 Work Breakdown Structure (WBS)")
    body(doc, "The Work Breakdown Structure breaks down our entire project into six major phases, each with specific tasks and sub-deliverables. This is basically our roadmap — it tells us what needs to happen, in what order, and roughly when it should be completed.")
    img(doc, "1.10 Work breakdown structure_diagram_1.svg", "Figure 1.2: Work Breakdown Structure — Mental Health Support Platform", 6.5)

    body(doc, "Here is a more detailed look at each phase:", bold=True)
    table(doc, ["Phase / Timeline", "Key Tasks and Deliverables"], [
        ["Phase 1: Discovery\n(Week 1–2)", "Problem statement definition, market research on existing mental health apps, user persona creation, competitor analysis, and establishing the project vision. This is where we figured out exactly what problem we are solving."],
        ["Phase 2: Requirements\n(Week 3–4)", "Documenting all 75 functional requirements (FR1 through FR75), defining non-functional requirements for performance, security, usability, reliability, and privacy. Also includes creating use case diagrams and user stories. The SRS document comes out of this phase."],
        ["Phase 3: Design\n(Week 5–6)", "Database schema design with the full ERD (15 entities), API endpoint specifications, wireframe creation in Balsamiq Cloud, all Data Flow Diagrams using Gane & Sarson notation, sequence diagrams, and the UML class diagram. This phase produces all the design documentation."],
        ["Phase 4: Development\n(Week 7–9)", "The actual coding happens here. Backend development using Express.js — all REST APIs, database implementation, authentication module. Frontend development using React.js — dashboard, mood tracker, messaging UI, appointment booking, and admin panel. AI chatbot integration with OpenAI, real-time messaging with Socket.io."],
        ["Phase 5: Testing\n(Week 10–11)", "Comprehensive testing at all levels — unit testing with Jest, integration testing with Supertest, end-to-end testing with Cypress, security testing using OWASP guidelines, and user acceptance testing with stakeholders."],
        ["Phase 6: Deployment\n(Week 11–12)", "Cloud deployment on Render or AWS, CI/CD pipeline configuration with GitHub Actions, final documentation, and the live demonstration presentation."],
    ], col_widths=[1.5, 5.2])
    add_footer_text(doc, "Mental Health Support Platform — Section 1.10 Work Breakdown Structure")
    save(doc, "10_Work_Breakdown_Structure.docx")


def gen_11_gantt():
    doc = new_doc()
    add_section_header(doc, "Section 1.11", "Gantt Chart", "Project Timeline Week by Week")
    heading(doc, "1.11 Gantt Chart / Project Timeline")
    body(doc, "The Gantt chart below provides a week-by-week timeline of our project. We have tried to be realistic with the scheduling — giving enough time for each phase while also accounting for the fact that things do not always go according to plan.")
    table(doc, ["Period", "Phase", "Key Activities", "Status"], [
        ["Week 1–2", "Discovery & Planning", "Problem research, user personas, project vision, team formation", "Completed ✓"],
        ["Week 3", "Requirements Analysis", "SRS document, all 75 functional requirements, non-functional requirements", "Completed ✓"],
        ["Week 4", "Architecture Design", "Three-tier architecture, tech stack finalisation, API specification", "Completed ✓"],
        ["Week 5", "Database & API Design", "ERD with 15 entities, DFD diagrams (Level 0, 1, 2), data dictionary", "Completed ✓"],
        ["Week 6", "UI/UX Design", "Wireframes in Balsamiq (6 screens), design review, accessibility check", "Completed ✓"],
        ["Week 7", "Backend Dev (Part 1)", "Auth module (JWT + OTP + bcrypt), user CRUD APIs, appointment APIs", "In Progress"],
        ["Week 8", "Backend Dev (Part 2)", "Mood tracking APIs, messaging with AES-256, chatbot integration, emergency APIs", "In Progress"],
        ["Week 9", "Frontend Development", "React components, dashboard, mood tracker UI, messaging UI, admin panel", "Planned"],
        ["Week 10", "Integration & Testing", "Frontend-backend integration, Socket.io real-time features, E2E testing", "Planned"],
        ["Week 11", "Final Testing & Deploy", "Full testing suite execution, security audit, cloud deployment, CI/CD setup", "Planned"],
        ["Week 12", "Demo & Handover", "Final presentation, live demonstration, documentation handover", "Planned"],
    ], col_widths=[0.9, 1.5, 3.0, 1.0])
    body(doc, "As you can see, we front-loaded the design and planning phases so that by the time we start coding in Week 7, we have a very clear picture of what needs to be built. This approach has saved us a lot of time and prevented the kind of confusion that happens when teams jump straight into coding without proper planning.")
    add_footer_text(doc, "Mental Health Support Platform — Section 1.11 Gantt Chart")
    save(doc, "11_Gantt_Chart.docx")


def gen_12_functional_req():
    doc = new_doc()
    add_section_header(doc, "Section 2.1", "Functional Requirements", "All 75 Requirements That the System Must Fulfil")
    heading(doc, "2.1 Functional Requirements")
    body(doc, "We have defined a total of 75 functional requirements for this platform, and each one of them is specific, measurable, and testable. We did not just write vague statements like 'the system should be secure' — instead, every requirement has concrete parameters that can be verified during testing.")
    body(doc, "The requirements are organised by module, and each one specifies what CRUD operation it involves. This makes it easier for the development team to know exactly what API endpoint needs to be built for each requirement.")

    fr_data = [
        ["FR1", "Register new user account with unique email/SIM validation", "Create", "Authentication"],
        ["FR2", "Verify accounts via OTP within a 5-second delivery window", "Create", "Authentication"],
        ["FR3", "Lock user account after 5 consecutive failed login attempts", "Update", "Authentication"],
        ["FR4", "Record user login timestamp and IP address for audit", "Create", "Audit"],
        ["FR5", "Process password reset with a 10-minute token validity", "Update", "Authentication"],
        ["FR6", "Update profile fields with strict input validation rules", "Update", "Profile"],
        ["FR7", "Upload and store profile image (maximum size 2MB)", "Create", "Profile"],
        ["FR8", "User login and logout with session management", "Read", "Authentication"],
        ["FR9", "Schedule appointment with time-slot conflict detection", "Create", "Appointment"],
        ["FR10", "Reject booking if the selected timeslot is unavailable", "Read", "Appointment"],
        ["FR11", "Reschedule appointments up to 2 hours before the session", "Update", "Appointment"],
        ["FR12", "Mark appointment as missed after a 15-minute no-show delay", "Update", "Appointment"],
        ["FR13", "Dispatch booking confirmation notification within 1 minute", "Create", "Notification"],
        ["FR14", "Approve pending appointments within a 10-minute window", "Update", "Appointment"],
        ["FR15", "Auto-cancel pending bookings after 30 minutes of inactivity", "Delete", "Appointment"],
        ["FR16", "Record clinical session notes after each consultation", "Create", "Clinical"],
        ["FR17", "Restrict session note access based on patient consent flags", "Read", "Security"],
        ["FR18", "View comprehensive patient history (authorised clinicians only)", "Read", "Clinical"],
        ["FR19", "Retain patient records independently of clinician employment", "Read", "Compliance"],
        ["FR20", "Reassign patients to another clinician upon departure", "Update", "Clinical"],
        ["FR21", "Request read-only access to previous medical records", "Read", "Clinical"],
        ["FR22", "Log every single medical record access and viewing action", "Create", "Audit"],
        ["FR23", "Encrypt all chat messages using AES-256 before storage", "Create", "Security"],
        ["FR24", "Deliver chat messages with less than 2 seconds end-to-end latency", "Create", "Messaging"],
        ["FR25", "Store chat transcripts permanently in the database", "Create", "Database"],
        ["FR26", "Delete chat messages locally using soft delete approach", "Delete", "Messaging"],
        ["FR27", "Submit daily mood entry with an exact timestamp", "Create", "Mood"],
        ["FR28", "Prevent duplicate mood entries within a 1-hour timeframe", "Create", "Mood"],
        ["FR29", "Generate automated weekly mood summary reports", "Read", "Mood"],
        ["FR30", "Trigger risk alert if mood score is below 2 for 3 consecutive days", "Create", "AI"],
        ["FR31", "Recommend relevant resources based on mood patterns", "Read", "AI"],
        ["FR32", "Respond to user queries via the AI Chatbot within 3 seconds", "Create", "AI Chatbot"],
        ["FR33", "Escalate chatbot session to human if confidence is below 70%", "Update", "AI Chatbot"],
        ["FR34", "Record all chatbot conversation transcripts for review", "Create", "AI Chatbot"],
        ["FR35", "Admin can review chatbot conversation logs via dashboard", "Read", "Admin"],
        ["FR36", "Identify users who have been inactive for 7 consecutive days", "Read", "AI"],
        ["FR37", "Dispatch automated wellness check notifications for inactive users", "Create", "Notification"],
        ["FR38", "Activate emergency button with less than 1 second response time", "Create", "Emergency"],
        ["FR39", "Display crisis hotline numbers immediately upon emergency trigger", "Read", "Emergency"],
        ["FR40", "Log all emergency button activations for audit purposes", "Create", "Audit"],
        ["FR41", "Admin can add mental health resources via admin panel", "Create", "Content"],
        ["FR42", "Validate all resource URLs before publishing them", "Create", "Content"],
        ["FR43", "Track and aggregate resource view counts for analytics", "Read", "Analytics"],
        ["FR44", "Generate monthly platform analytics reports automatically", "Read", "Analytics"],
        ["FR45", "Execute automated daily database backups at midnight", "Create", "System"],
        ["FR46", "Restore database from backup within a 1-hour recovery window", "Read", "System"],
        ["FR47", "Log all system errors and exceptions automatically", "Create", "Monitoring"],
        ["FR48", "Dispatch instant alerts to admin for critical system errors", "Create", "Monitoring"],
        ["FR49", "Allow users to select their preferred interface language", "Update", "UI"],
        ["FR50", "Load main dashboard within 3 seconds maximum", "Read", "Performance"],
        ["FR51", "Allow clinician to upload clinical documents and worksheets", "Create", "Clinical"],
        ["FR52", "Validate document uploads to ensure PDF format only", "Create", "Clinical"],
        ["FR53", "Allow patients to download their personal medical records", "Read", "Compliance"],
        ["FR54", "Allow users to revoke data sharing consent instantaneously", "Update", "Security"],
        ["FR55", "Restrict all access immediately upon consent revocation", "Update", "Security"],
        ["FR56", "Log all modifications to user consent settings for audit", "Create", "Audit"],
        ["FR57", "Integrate external SMS API (Twilio) for notifications", "Create", "Integration"],
        ["FR58", "Retry failed external API calls up to 3 times automatically", "Update", "Integration"],
        ["FR59", "Queue notifications during service downtimes for later delivery", "Create", "Integration"],
        ["FR60", "Admin can deactivate user accounts via the admin panel", "Update", "Admin"],
        ["FR61", "Track active session duration per user for analytics", "Read", "Analytics"],
        ["FR62", "Auto-logout users after 15 minutes of inactivity", "Update", "Security"],
        ["FR63", "Validate all user inputs to prevent SQL injection attacks", "Create", "Security"],
        ["FR64", "Generate complete audit trails for all database transactions", "Create", "Audit"],
        ["FR65", "Assign and enforce role-based access control (RBAC)", "Update", "Security"],
        ["FR66", "Deny and log all unauthorised API requests", "Read", "Security"],
        ["FR67", "Synchronise appointments with Google Calendar API", "Create", "Integration"],
        ["FR68", "Dispatch follow-up notifications for missed appointments", "Create", "Notification"],
        ["FR69", "Allow rebooking directly from missed session alert", "Create", "Appointment"],
        ["FR70", "Track clinician workload metrics for management", "Read", "Analytics"],
        ["FR71", "Generate regulatory compliance reports on demand", "Read", "Compliance"],
        ["FR72", "Log complete history of all data modifications", "Create", "Audit"],
        ["FR73", "Restrict file uploads exceeding 5MB in size", "Create", "System"],
        ["FR74", "Export analytics reports to CSV format", "Read", "Admin"],
        ["FR75", "Perform scheduled monthly system maintenance", "Read", "System"],
    ]
    table(doc, ["FR#", "Requirement Description", "CRUD", "Module"], fr_data, col_widths=[0.5, 3.8, 0.6, 1.0])
    add_footer_text(doc, "Mental Health Support Platform — Section 2.1 Functional Requirements")
    save(doc, "12_Functional_Requirements.docx")


def gen_13_nonfunctional_req():
    doc = new_doc()
    add_section_header(doc, "Section 2.2", "Non-Functional Requirements", "Quality Attributes the System Must Maintain")
    heading(doc, "2.2 Non-Functional Requirements")
    body(doc, "Functional requirements tell us what the system should do. Non-functional requirements tell us how well it should do it. These are equally important — a system that works but is slow, insecure, or difficult to use is practically useless. We have categorised our NFRs into six groups:")

    heading(doc, "Performance Requirements", level=2)
    body(doc, "Speed matters. Nobody is going to use a mental health app that takes forever to load, especially when they are already feeling low and just want quick access.")
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRP01", "Every page must load within 3 seconds on a standard broadband connection", "Performance"],
        ["NFRP02", "API responses must fall within 500ms at the 95th percentile — meaning 95 out of 100 requests should be this fast", "Performance"],
        ["NFRP03", "The system must support at least 50 concurrent users without any degradation in performance", "Scalability"],
        ["NFRP04", "Database queries must execute within 200ms — anything slower and the user will start noticing lag", "Performance"],
    ], col_widths=[1.0, 4.5, 1.2])

    heading(doc, "Security Requirements", level=2)
    body(doc, "When you are dealing with mental health data, security is not optional — it is absolutely critical. A data breach involving someone's mood entries or therapy notes would be devastating.")
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRS01", "All passwords must be hashed using bcrypt with a minimum of 10 salt rounds — no plaintext passwords anywhere", "Security"],
        ["NFRS02", "JWT authentication tokens must expire after 1 hour, forcing re-authentication", "Security"],
        ["NFRS03", "All client-server communication must use HTTPS with TLS 1.2 or higher — no exceptions", "Security"],
        ["NFRS04", "All sensitive data at rest must be encrypted using AES-256 standard", "Security"],
        ["NFRS05", "All user inputs must be sanitised to prevent SQL injection and Cross-Site Scripting (XSS) attacks", "Security"],
        ["NFRS06", "Rate limiting must be enforced at 100 requests per IP per minute, returning HTTP 429 for violations", "Security"],
    ], col_widths=[1.0, 4.5, 1.2])

    heading(doc, "Usability Requirements", level=2)
    body(doc, "The platform should be so intuitive that a new user can figure it out within minutes. We are building this for people who might already be stressed — the last thing they need is a confusing interface.")
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRU01", "A first-time user should be able to complete registration and start using the platform within 5 minutes", "Usability"],
        ["NFRU02", "The interface must be fully responsive across devices from 320px mobile to 1920px desktop viewport", "Usability"],
        ["NFRU03", "The platform must comply with WCAG 2.1 Level AA accessibility standards", "Accessibility"],
        ["NFRU04", "All error messages must be descriptive and actionable — telling the user what went wrong and how to fix it", "Usability"],
    ], col_widths=[1.0, 4.5, 1.2])

    heading(doc, "Reliability Requirements", level=2)
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRR01", "The system must maintain 99% uptime availability", "Reliability"],
        ["NFRR02", "Non-critical features should degrade gracefully without crashing the entire system", "Reliability"],
        ["NFRR03", "Automated daily backups with 7-day retention must be in place for disaster recovery", "Reliability"],
    ], col_widths=[1.0, 4.5, 1.2])

    heading(doc, "Maintainability Requirements", level=2)
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRM01", "ESLint must be enforced across the entire codebase with consistent coding standards", "Maintainability"],
        ["NFRM02", "The architecture must follow modular design with clear separation of concerns", "Maintainability"],
        ["NFRM03", "The system design must support 10x user growth without requiring a complete rewrite", "Scalability"],
    ], col_widths=[1.0, 4.5, 1.2])

    heading(doc, "Privacy and Compliance Requirements", level=2)
    table(doc, ["ID", "Requirement", "Category"], [
        ["NFRC01", "Privacy by design principles must be followed throughout the entire development process", "Privacy"],
        ["NFRC02", "Granular consent management must allow users to control exactly what data is shared and with whom", "Compliance"],
        ["NFRC03", "All analytics data must be anonymised so that individual users cannot be identified from reports", "Privacy"],
        ["NFRC04", "Users must have the right to request complete deletion of their data within 30 days", "Compliance"],
    ], col_widths=[1.0, 4.5, 1.2])
    add_footer_text(doc, "Mental Health Support Platform — Section 2.2 Non-Functional Requirements")
    save(doc, "13_NonFunctional_Requirements.docx")


def gen_14_architecture():
    doc = new_doc()
    add_section_header(doc, "Section 3.1", "Software & Hardware Architecture", "The Three-Tier System Design")
    heading(doc, "3.1 Software and Hardware Architecture")
    body(doc, "We went with a three-tier architecture for this platform — Presentation Tier, Application Tier, and Data Tier. This is a well-established pattern that provides clear separation of concerns, and it makes the system much easier to maintain, test, and scale independently.")
    img(doc, "3. System Architecture_diagram_1.svg", "Figure 3.1: Three-Tier System Architecture Diagram", 6.3)

    heading(doc, "Presentation Tier (Client Side)", level=2)
    body(doc, "This is what the user actually sees and interacts with. We are building it as a Single Page Application (SPA) using React.js 18+ with Vite as the build tool. TypeScript is being used for type safety — it catches a lot of bugs at compile time that would otherwise slip through. Socket.io is handling the client side of real-time features like chat and notifications. Chart.js is powering all the mood analytics visualisations. And Axios is our HTTP client for communicating with the backend, configured with a JWT interceptor that automatically attaches the auth token to every request.")

    heading(doc, "Application Tier (Server Side)", level=2)
    body(doc, "This is the brain of the operation. Express.js running on Node.js handles all API requests through a middleware chain that takes care of authentication, validation, error handling, and logging. The authentication module combines JWT tokens, OTP verification, bcrypt password hashing, and RBAC to create a robust security layer. The Socket.io server manages the real-time chat engine with AES-256 encryption and user presence tracking. External integrations connect to OpenAI for the chatbot, Twilio for SMS, SendGrid for emails, and Google Calendar for appointment sync.")

    heading(doc, "Data Tier (Storage)", level=2)
    body(doc, "PostgreSQL serves as our primary relational database, and honestly, it is perfect for this kind of application where data integrity and ACID compliance matter a lot. The database stores everything — user accounts, appointments, clinical records, encrypted messages, mood entries, and more. We also have file storage for profile images and clinical documents, audit logs for tracking every access and modification, and a cache/session layer for JWT tokens and API response caching.")
    add_footer_text(doc, "Mental Health Support Platform — Section 3.1 Software & Hardware Architecture")
    save(doc, "14_System_Architecture.docx")


def gen_15_tech_stack():
    doc = new_doc()
    add_section_header(doc, "Section 3.2", "Tech Stack", "Technologies and Tools We Are Using")
    heading(doc, "3.2 Tech Stack (Detailed)")
    body(doc, "Choosing the right technology stack was one of the most important early decisions we made. We wanted technologies that are modern, well-documented, have large community support, and most importantly — ones that the team is either already comfortable with or can pick up quickly. Here is the complete breakdown:")
    table(doc, ["Component", "Technology", "Why We Chose It"], [
        ["Frontend Framework", "React.js 18+ with Vite", "Component-based architecture makes it easy to build reusable UI elements. Vite gives us blazing fast hot module replacement during development."],
        ["Type Safety", "TypeScript", "Catches type-related bugs at compile time rather than runtime. Makes the codebase much more maintainable as it grows."],
        ["State Management", "React Context + React Query", "React Context handles local UI state, while React Query manages server state caching and synchronisation. A clean separation."],
        ["UI Styling", "Tailwind CSS / Material UI", "Tailwind gives us utility-first CSS for quick styling, and Material UI provides pre-built accessible components."],
        ["Charts", "Chart.js / Recharts", "For mood trend visualisation and analytics dashboards. Both libraries are lightweight and highly customisable."],
        ["Real-time", "Socket.io (Client + Server)", "The industry standard for WebSocket-based real-time communication. Handles chat and notifications seamlessly."],
        ["HTTP Client", "Axios", "Superior to fetch API — supports interceptors, request cancellation, and automatic JSON parsing."],
        ["Backend Runtime", "Node.js 18+", "Non-blocking I/O is perfect for handling multiple concurrent connections. Same language (JavaScript) on both ends."],
        ["API Framework", "Express.js", "Minimalist, flexible, and has the biggest middleware ecosystem in the Node.js world."],
        ["Authentication", "JWT + bcrypt + OTP", "JWT for stateless auth, bcrypt for secure password hashing, OTP for two-factor verification."],
        ["Encryption", "AES-256", "Military-grade encryption standard for protecting sensitive messages and clinical data."],
        ["Database", "PostgreSQL 15+", "Rock-solid relational database with ACID compliance, JSON support, and excellent performance."],
        ["ORM", "Prisma / Sequelize", "Provides database abstraction layer, migration management, and type-safe queries."],
        ["AI / NLP", "OpenAI API (GPT-3.5/4)", "Powers our chatbot with state-of-the-art natural language processing capabilities."],
        ["SMS", "Twilio", "Reliable SMS gateway for sending OTP codes and appointment notifications."],
        ["Email", "SendGrid", "Handles transactional emails — booking confirmations, password resets, wellness reminders."],
        ["Calendar", "Google Calendar API", "Synchronises appointments so users can see them alongside their personal events."],
        ["Unit Testing", "Jest", "Fast, well-documented testing framework with good mocking capabilities."],
        ["E2E Testing", "Cypress", "Simulates real user interactions in a browser — the best way to test complete workflows."],
        ["Version Control", "Git / GitHub", "Industry standard. We use feature branches, pull requests, and code reviews."],
        ["CI/CD", "GitHub Actions", "Automates our testing and deployment pipeline directly from the repository."],
        ["Hosting", "Render / AWS", "Render for simplicity, AWS as a backup option if we need more control."],
        ["Wireframes", "Balsamiq Cloud", "Quick low-fidelity wireframes that focus on layout and functionality, not pretty colours."],
    ], col_widths=[1.3, 1.8, 3.5])
    add_footer_text(doc, "Mental Health Support Platform — Section 3.2 Tech Stack")
    save(doc, "15_Tech_Stack.docx")


def gen_16_wireframes():
    doc = new_doc()
    add_section_header(doc, "Section 4.1", "User Interface Design / Wireframes", "How the Platform Will Look")
    heading(doc, "4.1 User Interface Design — Wireframes")
    body(doc, "We designed these wireframes using Balsamiq Cloud before writing a single line of frontend code. The idea was to get everyone on the same page about what the interface should look like and how users would navigate through the platform. These are intentionally low-fidelity — we focused on layout, information hierarchy, and user flow rather than colours and fonts.")

    wf_data = [
        ("4.1 User interface Design Wireframes_diagram_1.svg", "Figure 4.1: Login / Registration Page", "This is the first screen users will see. It has a clean split layout — login form on one side, registration option on the other. We have included fields for email, password, and a prominent OTP verification step. The 'Forgot Password' link is easily visible because users tend to forget passwords quite often."),
        ("4.1 User interface Design Wireframes_diagram_2.svg", "Figure 4.2: User Dashboard", "The dashboard is the central hub of the platform. It shows a mood summary chart at the top, upcoming appointments, recent chat messages, and quick-access links to key features. We designed it to give users a complete overview of their mental health journey at just one glance."),
        ("4.1 User interface Design Wireframes_diagram_3.svg", "Figure 4.3: Mood Tracker Page", "Users can log their daily mood on a 1–10 scale with an optional notes field. The page also shows a trend graph of past mood entries so users can spot patterns. There is a colour-coded indicator — green for good days, yellow for neutral, and red for concerning entries."),
        ("4.1 User interface Design Wireframes_diagram_4.svg", "Figure 4.4: Secure Messaging Interface", "The messaging interface looks like a standard chat application — familiar and intuitive. It has a contact list on the left and the conversation thread on the right. A small lock icon next to each message indicates that it is encrypted. Users can see read receipts and typing indicators."),
        ("4.1 User interface Design Wireframes_diagram_5.svg", "Figure 4.5: Appointments Page", "Shows available time slots in a calendar view. Users can select a clinician, pick a date, and choose from available slots. Booked appointments appear in a list below with options to reschedule or cancel. The Google Calendar sync icon is visible for connected users."),
        ("4.1 User interface Design Wireframes_diagram_6.svg", "Figure 4.6: Admin Dashboard Panel", "The admin panel gives administrators a bird's-eye view of the entire platform — user statistics, active sessions, system health indicators, recent audit logs, and quick actions like user management and content moderation. It is designed for efficiency."),
    ]
    for svg_name, caption, desc in wf_data:
        body(doc, desc, size=10)
        img(doc, svg_name, caption, 5.5)

    add_footer_text(doc, "Mental Health Support Platform — Section 4.1 User Interface Design")
    save(doc, "16_Wireframes.docx")


def gen_17_use_case():
    doc = new_doc()
    add_section_header(doc, "Section 5", "Use Case Diagrams", "How Different Users Interact with the System")
    heading(doc, "5. Use Case Diagrams")
    body(doc, "Use case diagrams are honestly one of the most useful diagrams for understanding a system because they show you, at a glance, what each type of user can do. We have three primary actors in our system — Patient/User, Clinician/Doctor, and Administrator — and each of them has a different set of capabilities.")

    heading(doc, "5.1 Patient / User — Use Case Diagram", level=2)
    body(doc, "The patient is the primary user of the platform. They have the most use cases because the entire system is built around supporting their mental health journey.")
    img(doc, "Patient  User — Use Case Diagram.svg", "Figure 5.1: Patient / User — Use Case Diagram", 6.0)
    body(doc, "Key use cases for the patient include:", bold=True)
    for uc in [
        "Account Management — Register (FR1), Login/Logout (FR8), Verify OTP (FR2), Reset Password (FR5)",
        "Profile — Update personal details (FR6), Upload profile image (FR7)",
        "Mood Tracking — Submit daily mood entry (FR27), View weekly mood reports (FR29)",
        "Appointments — Book appointments (FR9), Reschedule up to 2 hours before (FR11)",
        "Communication — Send encrypted chat messages (FR23), Interact with AI Chatbot (FR32)",
        "Support — Access resource library (FR41), Trigger emergency button (FR38)",
        "Privacy — Manage consent settings (FR54), Download personal records (FR53), Select language (FR49)",
    ]:
        bullet(doc, uc)

    heading(doc, "5.2 Clinician / Doctor — Use Case Diagram", level=2)
    body(doc, "The clinician interacts with the platform primarily to manage patient relationships, appointments, and clinical records.")
    img(doc, "Clinician  Doctor — Use Case Diagram.svg", "Figure 5.2: Clinician / Doctor — Use Case Diagram", 6.0)
    body(doc, "Key use cases for the clinician:", bold=True)
    for uc in [
        "Patient Management — View patient history (FR18), Write session notes (FR16), Request medical records (FR21)",
        "Appointments — Approve pending appointments (FR14), View schedule synced with Google Calendar (FR67)",
        "Communication — Chat with patients securely (FR24), Receive chatbot escalations (FR33)",
        "Clinical — Upload clinical documents (FR51), Reassign patients (FR20), Track workload metrics (FR70)",
    ]:
        bullet(doc, uc)

    heading(doc, "5.3 Administrator — Use Case Diagram", level=2)
    body(doc, "The administrator manages the overall platform operations, user accounts, content, and system health.")
    img(doc, "Administrator — Use Case Diagram.svg", "Figure 5.3: Administrator — Use Case Diagram", 6.0)
    body(doc, "Key use cases for the administrator:", bold=True)
    for uc in [
        "User Management — Manage user accounts (FR60), Assign RBAC roles (FR65), Deactivate accounts",
        "Content — Add and manage mental health resources (FR41), Validate resource URLs (FR42)",
        "Analytics — View platform analytics dashboard (FR44), Generate compliance reports (FR71), Export to CSV (FR74)",
        "System — Review chatbot logs (FR35), Monitor system health (FR47), Create/Restore backups (FR45–46), View audit trails (FR64)",
    ]:
        bullet(doc, uc)
    add_footer_text(doc, "Mental Health Support Platform — Section 5 Use Case Diagrams")
    save(doc, "17_Use_Case_Diagrams.docx")


def gen_18_context_diagram():
    doc = new_doc()
    add_section_header(doc, "Section 6", "Context Diagram", "Level 0 — The Big Picture View")
    heading(doc, "6. Context Diagrams (Gane & Sarson Method)")
    body(doc, "The Level 0 Context Diagram is the highest-level view of our system. It shows the Mental Health Support Platform as a single central process and all the external entities that interact with it. We have used the Gane & Sarson notation where processes are shown as rounded rectangles and external entities as rectangular boxes.")
    body(doc, "This diagram is really useful for getting a quick understanding of the system boundaries — what is inside the system and what is outside. Every arrow represents a data flow between an external entity and the platform.")
    img(doc, "Level 0 — Context Diagram (Gane & Sarson).svg", "Figure 6.1: Level 0 — Context Diagram (Gane & Sarson)", 6.3)

    heading(doc, "External Entities", level=2)
    body(doc, "Our system interacts with seven external entities:")
    entities = [
        ("E1: Patient/User", "The primary user of the platform. They send registration data, mood entries, appointment requests, chat messages, and emergency triggers. In return, they receive dashboard data, mood reports, appointment confirmations, encrypted messages, and crisis information."),
        ("E2: Clinician/Doctor", "Healthcare professionals who interact with the system to manage appointments, write clinical notes, chat with patients, and view patient history. They receive patient data, appointment schedules, and escalated chatbot sessions."),
        ("E3: Administrator", "System administrators who manage users, resources, analytics, and system configuration. They receive analytics data, audit reports, and system health information."),
        ("E4: SMS Gateway (Twilio)", "External service that receives SMS notification requests from the platform and delivers OTP codes, appointment reminders, and wellness check messages to users' phones."),
        ("E5: Google Calendar", "External calendar service that receives appointment data for synchronisation, allowing users and clinicians to see their appointments alongside personal calendar events."),
        ("E6: Email Service (SendGrid)", "External email service that handles transactional emails — booking confirmations, password reset links, weekly mood summaries, and follow-up notifications."),
        ("E7: Crisis Hotline", "External crisis support service that provides hotline numbers and support data which the platform displays when a user activates the emergency button."),
    ]
    for name, desc in entities:
        body(doc, name, bold=True, size=11)
        body(doc, desc, size=10)
    add_footer_text(doc, "Mental Health Support Platform — Section 6 Context Diagram")
    save(doc, "18_Context_Diagram.docx")


def gen_19_dfd_level1():
    doc = new_doc()
    add_section_header(doc, "Section 7.1", "DFD Level 1", "System-Wide Data Flow Decomposition")
    heading(doc, "7.1 Data Flow Diagram — Level 1")
    body(doc, "The Level 1 DFD breaks down the single central process from the Context Diagram into 15 major sub-processes. Each sub-process handles a specific area of functionality and interacts with relevant data stores and external entities. This gives us a much clearer picture of how data actually flows through the system.")
    img(doc, "Level 1 — System Data Flow Diagram.svg", "Figure 7.1: Level 1 — System Data Flow Diagram", 6.3)

    heading(doc, "Process Descriptions", level=2)
    table(doc, ["Process", "Name", "Functional Requirements"], [
        ["1.0", "Authentication & Account Management", "FR1–5, FR8"],
        ["2.0", "Profile Management", "FR6–7"],
        ["3.0", "Appointment Management", "FR9–15, FR67–69"],
        ["4.0", "Clinical Records Management", "FR16–22, FR51–53"],
        ["5.0", "Messaging System", "FR23–26"],
        ["6.0", "Mood Tracking & Analysis", "FR27–31"],
        ["7.0", "AI Chatbot System", "FR32–36"],
        ["8.0", "Emergency Support", "FR38–40"],
        ["9.0", "Notification Engine", "FR13, FR37, FR57–59, FR68"],
        ["10.0", "Content & Resource Management", "FR41–43"],
        ["11.0", "Analytics & Reporting", "FR44, FR61, FR70–72, FR74"],
        ["12.0", "Administration & Access Control", "FR60, FR62–66"],
        ["13.0", "System Maintenance & Backup", "FR45–48, FR73, FR75"],
        ["14.0", "Privacy & Compliance", "FR54–56"],
        ["15.0", "UI & Localisation", "FR49–50"],
    ], col_widths=[0.8, 2.8, 2.5])

    heading(doc, "Data Stores", level=2)
    body(doc, "The system uses 15 data stores, each dedicated to a specific type of data:")
    table(doc, ["ID", "Data Store Name", "Purpose"], [
        ["D1", "Users & Credentials", "Stores user accounts, hashed passwords, verification status, and login history"],
        ["D2", "User Profiles", "Stores display names, avatar images, language preferences, and profile metadata"],
        ["D3", "Appointments", "Stores all appointment records with status tracking, clinician assignments, and calendar event IDs"],
        ["D4", "Clinical Records", "Stores session notes, medical histories, clinical documents, and retention policies"],
        ["D5", "Chat Messages", "Stores AES-256 encrypted message bodies, delivery status, and soft delete flags"],
        ["D6", "Mood Entries", "Stores daily mood scores (1–10), notes, timestamps, and entry dates"],
        ["D7", "Chatbot Transcripts", "Stores all AI chatbot conversations with confidence scores and escalation flags"],
        ["D8", "Emergency Logs", "Stores emergency button activation records with timestamps and resolution status"],
        ["D9", "Notification Queue", "Stores pending, sent, failed, and queued notifications across all channels"],
        ["D10", "Resource Library", "Stores curated mental health articles, URLs, categories, and view counts"],
        ["D11", "Analytics Data", "Stores aggregated platform metrics, session durations, and report data"],
        ["D12", "Audit Logs", "Stores comprehensive audit trails for every system action and data access event"],
        ["D13", "System Backups", "Stores automated database backup records and restoration history"],
        ["D14", "Consent Records", "Stores user consent preferences, data sharing flags, and consent change history"],
        ["D15", "Session Data", "Stores active session tokens, JWT data, and API response cache"],
    ], col_widths=[0.5, 1.5, 4.5])
    add_footer_text(doc, "Mental Health Support Platform — Section 7.1 DFD Level 1")
    save(doc, "19_DFD_Level_1.docx")


def gen_20_dfd_level2():
    doc = new_doc()
    add_section_header(doc, "Section 7.2", "DFD Level 2", "Detailed Process Decomposition")
    heading(doc, "7.2 Data Flow Diagrams — Level 2")
    body(doc, "Each of the 15 Level 1 processes is further decomposed into detailed sub-processes in the Level 2 DFDs. These diagrams show the internal logic, decision points, and data store interactions within each process. They are extremely useful for the development team because they essentially serve as a blueprint for how each module should be coded.")

    level2 = [
        ("Level 2.1 — Authentication & Account Management.svg", "Figure 7.2.1: Level 2.1 — Authentication & Account Management",
         "This diagram breaks down the authentication process into 7 sub-processes: User Registration (FR1) which validates email uniqueness and creates the account, OTP Verification (FR2) with a 5-second delivery window, Account Locking (FR3) after 5 failed attempts, Login Activity Logging (FR4) that records timestamps and IP addresses, Password Reset (FR5) with 10-minute token validity, Login/Logout (FR8) with JWT token management, and Account Deletion with data anonymisation."),
        ("Level 2.2 — Profile Management.svg", "Figure 7.2.2: Level 2.2 — Profile Management",
         "A simpler diagram with just 2 sub-processes: Update Profile Fields (FR6) which validates all input data before saving changes, and Upload Profile Image (FR7) which enforces the 2MB file size limit and validates image formats before storing."),
        ("Level 2.3 — Appointment Management.svg", "Figure 7.2.3: Level 2.3 — Appointment Management",
         "One of the more complex Level 2 diagrams with 9 sub-processes covering the complete appointment lifecycle: scheduling with conflict detection (FR9), slot availability checking (FR10), rescheduling with the 2-hour cutoff rule (FR11), missed marking after 15-minute delay (FR12), clinician approval within 10 minutes (FR14), auto-cancellation after 30 minutes of inactivity (FR15), Google Calendar synchronisation (FR67), follow-up notification dispatch (FR68), and rebooking from missed alerts (FR69)."),
        ("Level 2.4 — Clinical Records Management.svg", "Figure 7.2.4: Level 2.4 — Clinical Records Management",
         "Contains 10 sub-processes for managing sensitive clinical data: session note recording (FR16), consent-based access checking (FR17), patient history viewing for authorised clinicians (FR18), record retention independent of clinician employment (FR19), patient reassignment upon clinician departure (FR20), medical record access requests (FR21), comprehensive access logging (FR22), clinical document upload (FR51), PDF format validation (FR52), and secure record download (FR53)."),
        ("Level 2.5 — Messaging System.svg", "Figure 7.2.5: Level 2.5 — Messaging System",
         "Shows 4 sub-processes for the secure messaging module: message encryption using AES-256 before storage (FR23), real-time message delivery within 2 seconds via Socket.io (FR24), permanent chat transcript storage in the database (FR25), and local soft delete functionality that hides messages without purging them (FR26)."),
        ("Level 2.6 — Mood Tracking & Analysis.svg", "Figure 7.2.6: Level 2.6 — Mood Tracking & Analysis",
         "Decomposes into 5 sub-processes: daily mood entry submission with timestamp recording (FR27), duplicate prevention within a 1-hour window to avoid accidental double-entries (FR28), automated weekly mood report generation (FR29), risk alert triggering when the mood score drops below 2 for 3 consecutive days (FR30), and intelligent resource recommendation based on detected mood patterns (FR31)."),
        ("Level 2.7 — AI Chatbot System.svg", "Figure 7.2.7: Level 2.7 — AI Chatbot System",
         "Contains 5 sub-processes: user query processing through the OpenAI NLP engine with a 3-second response target (FR32), automatic escalation to a human clinician when AI confidence falls below 70% (FR33), conversation transcript recording for quality review (FR34), admin dashboard log review capability (FR35), and inactive user detection after 7 days of no activity to trigger wellness checks (FR36)."),
        ("Level 2.8 — Emergency Support.svg", "Figure 7.2.8: Level 2.8 — Emergency Support",
         "A focused diagram with 3 sub-processes: emergency button activation with strict less-than-1-second response time (FR38), immediate crisis hotline number display from the external crisis service (FR39), and comprehensive emergency activation logging for audit and follow-up purposes (FR40)."),
        ("Level 2.9 — Notification Engine.svg", "Figure 7.2.9: Level 2.9 — Notification Engine",
         "Shows 6 sub-processes for the multi-channel notification system: booking confirmation dispatch within 1 minute (FR13), automated wellness check notifications for inactive users (FR37), SMS and email integration through Twilio and SendGrid (FR57), failed API call retry up to 3 times (FR58), notification queuing during service downtimes (FR59), and missed appointment follow-up notifications (FR68)."),
        ("Level 2.10 — Content & Resource Management.svg", "Figure 7.2.10: Level 2.10 — Content & Resource Management",
         "Contains 3 sub-processes: resource addition through the admin panel interface (FR41), URL validation before any resource is published to ensure links are valid and safe (FR42), and view count tracking and aggregation for analytics purposes (FR43)."),
        ("Level 2.11 — Analytics & Reporting.svg", "Figure 7.2.11: Level 2.11 — Analytics & Reporting",
         "Decomposes into 6 sub-processes: monthly platform analytics report generation (FR44), active session duration tracking per user (FR61), clinician workload metrics evaluation (FR70), regulatory compliance report generation (FR71), complete data modification history logging (FR72), and CSV export functionality for reports (FR74)."),
        ("Level 2.12 — Administration & Access Control.svg", "Figure 7.2.12: Level 2.12 — Administration & Access Control",
         "Shows 6 sub-processes: user account deactivation via admin panel (FR60), auto-logout enforcement after 15 minutes of inactivity (FR62), input validation and sanitisation for SQL injection prevention (FR63), comprehensive audit trail generation for all database transactions (FR64), RBAC assignment and enforcement (FR65), and unauthorised API request denial with logging (FR66)."),
        ("Level 2.13 — System Maintenance & Backup.svg", "Figure 7.2.13: Level 2.13 — System Maintenance & Backup",
         "Contains 6 sub-processes: automated daily database backup execution at midnight (FR45), backup restoration within a 1-hour recovery window (FR46), real-time system error and exception logging (FR47), critical error alert dispatch to administrators (FR48), file upload size restriction enforcement at 5MB (FR73), and scheduled monthly system maintenance tasks (FR75)."),
        ("Level 2.14 — Privacy & Compliance.svg", "Figure 7.2.14: Level 2.14 — Privacy & Compliance",
         "A focused diagram with 3 sub-processes: instant consent revocation that takes effect immediately (FR54), immediate access restriction upon consent revocation so that no data can be viewed after the user withdraws consent (FR55), and comprehensive consent change logging that records every modification for regulatory compliance (FR56)."),
        ("Level 2.15 — UI & Localization.svg", "Figure 7.2.15: Level 2.15 — UI & Localisation",
         "The simplest Level 2 diagram with 2 sub-processes: interface language selection that allows users to choose their preferred language (FR49), and dashboard loading performance optimisation to ensure the main dashboard renders within 3 seconds (FR50)."),
    ]
    for svg_name, caption, desc in level2:
        body(doc, desc, size=10)
        img(doc, svg_name, caption, 6.0)
    add_footer_text(doc, "Mental Health Support Platform — Section 7.2 DFD Level 2")
    save(doc, "20_DFD_Level_2.docx")


def gen_21_sequence():
    doc = new_doc()
    add_section_header(doc, "Section 8", "Sequence Diagrams", "Step-by-Step Interaction Flows")
    heading(doc, "8. Sequence Diagrams")
    body(doc, "Sequence diagrams are brilliant for showing exactly what happens when a user performs a specific action — which components talk to each other, in what order, and what data gets passed around. We have created six sequence diagrams covering the most critical user flows in the system.")

    sd_data = [
        ("SD1 User Registration & Login.svg", "Figure 8.1: SD1 — User Registration & Login",
         "This is the very first interaction any user will have with our system. The flow goes like this — the patient fills in the registration form on the frontend, which sends the data to the API server. The server validates everything, creates a new record in the database, generates an OTP, and sends it via email. The patient enters the OTP, the server verifies it, marks the account as verified, generates a JWT token, and the frontend renders the dashboard. If the OTP is wrong or expired, the user gets a clear error. Simple, secure, and straightforward."),
        ("SD2 Appointment Booking Flow.svg", "Figure 8.2: SD2 — Appointment Booking Flow",
         "When a patient wants to book an appointment, they select a clinician and a date on the frontend. The API server checks available time slots against the database, returns the available ones, and the patient picks a slot. The server then checks for conflicts (because two patients should not be able to book the same slot), creates the appointment record, pushes an event to Google Calendar, and dispatches a confirmation email through SendGrid. The whole process takes just a few seconds."),
        ("SD3 Mood Entry & Risk Alert.svg", "Figure 8.3: SD3 — Mood Entry & Risk Alert",
         "The patient submits a mood entry with a score and optional notes. The server first checks if there is already an entry within the last 24 hours (to prevent duplicates), inserts the new entry, and then runs a risk assessment. If the mood score has been below 2 for three consecutive days, the system automatically creates a risk alert notification and flags it for clinician review. The frontend dashboard updates in real-time to reflect the new entry and any alerts."),
        ("SD4 AI Chatbot Interaction.svg", "Figure 8.4: SD4 — AI Chatbot Interaction",
         "The patient types a message to the chatbot. The frontend sends it to the API server, which forwards it to the OpenAI engine for NLP processing. The AI generates a response with a confidence score. If confidence is 70% or above, the response goes straight back to the patient within 3 seconds. If it drops below 70%, the system automatically escalates the conversation to a human clinician and notifies both parties. Every conversation is stored as a transcript for admin review."),
        ("SD5 Emergency Button Activation.svg", "Figure 8.5: SD5 — Emergency Button Activation",
         "This is a time-critical flow — when someone presses the emergency button, every millisecond counts. The frontend sends an activation request, the API server creates an emergency log entry, fetches crisis hotline data from the external crisis service, and returns it to the frontend — all within less than 1 second. The frontend immediately displays crisis contact numbers, helpline information, and coping resources. An audit trail entry is also created for follow-up by the care team."),
        ("SD6 Secure Messaging.svg", "Figure 8.6: SD6 — Secure Messaging",
         "When a patient sends a message to their clinician, the frontend sends the message content to the API server. The server encrypts it using AES-256 (so even if someone gets access to the database, they cannot read the messages), stores the encrypted version, and then pushes it via WebSocket to the clinician's client in real-time. Delivery takes less than 2 seconds. The clinician receives a notification, and the sender gets a delivery confirmation. If a user soft-deletes a message, it disappears from their view but the encrypted version remains in the database for record-keeping."),
    ]
    for svg_name, caption, desc in sd_data:
        body(doc, desc, size=10)
        img(doc, svg_name, caption, 5.5)
    add_footer_text(doc, "Mental Health Support Platform — Section 8 Sequence Diagrams")
    save(doc, "21_Sequence_Diagrams.docx")


def gen_22_erd():
    doc = new_doc()
    add_section_header(doc, "Section 9", "Entity Relationship Diagram\n& Data Dictionary", "Database Design")
    heading(doc, "9. Entity Relationship Diagram & Data Dictionary")
    body(doc, "The ERD represents the logical data model of our entire database. We have 15 entities (tables), each carefully normalised to avoid data redundancy while maintaining data integrity. Every relationship is defined with proper cardinality (one-to-many, many-to-many) and foreign key constraints.")
    img(doc, "Entity Relationship Diagram — Mental Health Support Platform.svg", "Figure 9.1: Entity Relationship Diagram — Mental Health Support Platform", 6.3)

    heading(doc, "9.1 Data Dictionary", level=2)
    body(doc, "The data dictionary below defines every column in every table. We have been very specific about data types, constraints, and what each column actually stores. This serves as the definitive reference for our backend developer when implementing the database schema.")

    dd = [
        ("Users", [
            ["user_id", "UUID", "PK", "Unique identifier for each user account"],
            ["email", "VARCHAR(255)", "UNIQUE, NOT NULL", "User's email address — must be unique across the system"],
            ["password_hash", "TEXT", "NOT NULL", "bcrypt hashed password with 10 salt rounds"],
            ["role", "ENUM", "NOT NULL", "One of: patient, clinician, or admin"],
            ["verified", "BOOLEAN", "DEFAULT false", "Whether the user has completed OTP verification"],
            ["locked", "BOOLEAN", "DEFAULT false", "Account locked after 5 failed login attempts"],
            ["mfa_enabled", "BOOLEAN", "DEFAULT false", "Whether multi-factor authentication is enabled"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the account was created"],
            ["last_login", "TIMESTAMP", "NULLABLE", "Timestamp of the most recent successful login"],
        ]),
        ("Profiles", [
            ["profile_id", "UUID", "PK", "Unique profile identifier"],
            ["user_id", "UUID", "FK → Users", "Links to the Users table"],
            ["display_name", "VARCHAR(100)", "NOT NULL", "Name shown on the user's profile and in chats"],
            ["avatar_url", "TEXT", "NULLABLE", "URL of uploaded profile image (maximum 2MB)"],
            ["language", "VARCHAR(10)", "DEFAULT 'en'", "User's preferred interface language"],
            ["updated_at", "TIMESTAMP", "DEFAULT NOW()", "When the profile was last modified"],
        ]),
        ("Appointments", [
            ["appt_id", "UUID", "PK", "Unique appointment identifier"],
            ["patient_id", "UUID", "FK → Users", "The patient who booked the appointment"],
            ["clinician_id", "UUID", "FK → Users", "The clinician assigned to the appointment"],
            ["date_time", "TIMESTAMP", "NOT NULL", "Scheduled date and time of the appointment"],
            ["status", "ENUM", "NOT NULL", "Current status: pending, confirmed, missed, or cancelled"],
            ["gcal_event_id", "TEXT", "NULLABLE", "Google Calendar event ID for synced appointments"],
            ["missed", "BOOLEAN", "DEFAULT false", "Flagged true if patient did not show up after 15 minutes"],
            ["auto_cancelled", "BOOLEAN", "DEFAULT false", "Flagged true if pending for more than 30 minutes"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the appointment was initially created"],
        ]),
        ("ClinicalRecords", [
            ["record_id", "UUID", "PK", "Unique clinical record identifier"],
            ["patient_id", "UUID", "FK → Users", "The patient this record belongs to"],
            ["clinician_id", "UUID", "FK → Users", "The clinician who created the record"],
            ["session_notes", "TEXT", "NOT NULL", "Clinical session notes written by the clinician"],
            ["record_type", "ENUM", "NOT NULL", "Type: consultation, follow-up, or emergency"],
            ["retention_date", "DATE", "NULLABLE", "Date until which the record must be retained"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the clinical record was created"],
        ]),
        ("ChatMessages", [
            ["message_id", "UUID", "PK", "Unique message identifier"],
            ["sender_id", "UUID", "FK → Users", "The user who sent the message"],
            ["receiver_id", "UUID", "FK → Users", "The intended recipient of the message"],
            ["encrypted_body", "TEXT", "NOT NULL", "AES-256 encrypted message content"],
            ["delivered", "BOOLEAN", "DEFAULT false", "Whether the message has been delivered to the recipient"],
            ["soft_deleted", "BOOLEAN", "DEFAULT false", "Soft delete flag — message hidden but not purged"],
        ]),
        ("MoodEntries", [
            ["mood_id", "UUID", "PK", "Unique mood entry identifier"],
            ["user_id", "UUID", "FK → Users", "The user who submitted this mood entry"],
            ["score", "INT", "CHECK (1–5)", "Mood score on a 1 to 5 scale"],
            ["notes", "TEXT", "NULLABLE", "Optional personal notes about how the user is feeling"],
            ["entry_date", "DATE", "NOT NULL", "The date of the mood entry"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the entry was recorded in the system"],
        ]),
        ("ChatbotTranscripts", [
            ["transcript_id", "UUID", "PK", "Unique transcript record identifier"],
            ["user_id", "UUID", "FK → Users", "The user who interacted with the chatbot"],
            ["query", "TEXT", "NOT NULL", "The original query text sent by the user"],
            ["response", "TEXT", "NOT NULL", "The AI-generated response text"],
            ["confidence", "FLOAT", "NOT NULL", "AI confidence score ranging from 0.0 to 1.0"],
            ["escalated", "BOOLEAN", "DEFAULT false", "Whether the conversation was escalated (confidence < 70%)"],
        ]),
        ("Notifications", [
            ["notif_id", "UUID", "PK", "Unique notification identifier"],
            ["user_id", "UUID", "FK → Users", "The user who will receive this notification"],
            ["type", "ENUM", "NOT NULL", "Notification type: booking, wellness, alert, or follow_up"],
            ["channel", "ENUM", "NOT NULL", "Delivery channel: email, sms, or in_app"],
            ["payload", "JSON", "NOT NULL", "The actual notification content and metadata"],
            ["status", "ENUM", "DEFAULT 'pending'", "Delivery status: pending, sent, failed, or queued"],
            ["retry_count", "INT", "DEFAULT 0", "Number of delivery retry attempts (maximum 3)"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the notification was created"],
        ]),
        ("EmergencyLogs", [
            ["log_id", "UUID", "PK", "Unique emergency log identifier"],
            ["user_id", "UUID", "FK → Users", "The user who activated the emergency button"],
            ["activated_at", "TIMESTAMP", "DEFAULT NOW()", "Exact timestamp of the emergency activation"],
            ["hotline_data", "JSON", "NULLABLE", "Crisis hotline information displayed to the user"],
            ["resolved", "BOOLEAN", "DEFAULT false", "Whether the emergency situation has been resolved"],
        ]),
        ("ConsentRecords", [
            ["consent_id", "UUID", "PK", "Unique consent record identifier"],
            ["user_id", "UUID", "FK → Users", "The user whose consent this record tracks"],
            ["consent_type", "ENUM", "NOT NULL", "Type: data_sharing, analytics, or marketing"],
            ["granted", "BOOLEAN", "NOT NULL", "Whether consent is currently granted or revoked"],
            ["changed_at", "TIMESTAMP", "DEFAULT NOW()", "When the consent setting was last changed"],
        ]),
        ("AuditLogs", [
            ["audit_id", "UUID", "PK", "Unique audit log entry identifier"],
            ["user_id", "UUID", "FK → Users", "The user who performed the action"],
            ["action", "VARCHAR(100)", "NOT NULL", "Description of the action performed"],
            ["entity_type", "VARCHAR(50)", "NOT NULL", "The type of entity that was affected"],
            ["entity_id", "UUID", "NULLABLE", "The specific entity record that was affected"],
            ["ip_address", "INET", "NULLABLE", "The client IP address from which the action was performed"],
            ["details", "JSON", "NULLABLE", "Additional context and metadata about the action"],
            ["created_at", "TIMESTAMP", "DEFAULT NOW()", "When the audit entry was recorded"],
        ]),
    ]
    for tname, cols in dd:
        body(doc, f"Table: {tname}", bold=True)
        table(doc, ["Column", "Type", "Constraints", "Description"], cols, col_widths=[1.3, 1.2, 1.4, 2.8])
    add_footer_text(doc, "Mental Health Support Platform — Section 9 ERD & Data Dictionary")
    save(doc, "22_ERD_Data_Dictionary.docx")


def gen_23_class_diagram():
    doc = new_doc()
    add_section_header(doc, "Section 10", "Class Diagram", "Object-Oriented System Model")
    heading(doc, "10. UML Class Diagram")
    body(doc, "The UML Class Diagram gives us an object-oriented view of the system. While the ERD shows us how data is stored, the class diagram shows us how the application logic is structured — what classes exist, what properties they have, what methods they expose, and how they relate to each other.")
    img(doc, "UML Class Diagram — Mental Health Support Platform.svg", "Figure 10.1: UML Class Diagram — Mental Health Support Platform", 6.3)

    heading(doc, "Class Descriptions", level=2)
    body(doc, "Our system consists of 12 primary classes. Each class encapsulates both data (attributes) and behaviour (methods) related to a specific domain concept:")
    classes = [
        ("User", "The core authentication class. Handles registration, login, OTP verification, password reset, and account deletion. Every other class in the system relates back to this one because everything revolves around authenticated users."),
        ("Profile", "Manages user profile information — display name, avatar image, language preference. Separated from User because profile data changes frequently while authentication data is relatively static."),
        ("Appointment", "The scheduling engine. Methods include scheduling with conflict detection, slot availability checking, rescheduling (with the 2-hour cutoff), approval by clinician, auto-cancellation, and marking appointments as missed."),
        ("ClinicalRecord", "Handles all medical record operations. Key methods include consent checking before allowing access, session note writing, document upload and download, patient history viewing, and patient reassignment between clinicians."),
        ("ChatMessage", "The secure messaging class. Responsible for AES-256 encryption before sending, real-time delivery through Socket.io, delivery confirmation, and soft-delete operations."),
        ("MoodEntry", "Manages the entire mood tracking lifecycle — entry submission with duplicate prevention, weekly report generation, risk alert triggering when patterns indicate concern, and resource recommendation based on mood trends."),
        ("Chatbot", "The AI interaction class. Processes user queries through the OpenAI NLP engine, handles response generation with confidence scoring, manages escalation to human clinicians, stores transcripts, and detects inactive users."),
        ("Emergency", "A focused class handling the critical emergency flow — button activation with sub-second response, crisis hotline data fetching from external services, and comprehensive activation logging."),
        ("Notification", "The multi-channel notification engine. Supports email (SendGrid), SMS (Twilio), and in-app notifications. Handles retry logic (up to 3 attempts), downtime queuing, and delivery status tracking."),
        ("Resource", "Manages the mental health resource library. Handles resource addition, URL validation before publishing, view count tracking, and content categorisation."),
        ("Analytics", "The reporting engine. Generates monthly platform reports, tracks clinician workload, produces compliance reports, logs data modifications, and exports data to CSV format."),
        ("Admin", "The administrative control class. Manages user accounts (activation/deactivation), RBAC role assignment, audit trail viewing, system backup creation and restoration, and platform-wide monitoring."),
    ]
    for name, desc in classes:
        body(doc, name, bold=True, size=11)
        body(doc, desc, size=10)
    add_footer_text(doc, "Mental Health Support Platform — Section 10 Class Diagram")
    save(doc, "23_Class_Diagram.docx")


def gen_24_dpia():
    doc = new_doc()
    add_section_header(doc, "Section 11", "Data Protection Impact\nAssessment (DPIA)", "Privacy Risk Analysis and Mitigation")
    heading(doc, "11. Data Protection Impact Assessment (DPIA)")
    body(doc, "Since we are dealing with mental health data — which is about as sensitive as personal data gets — we took the DPIA very seriously. This is not just a compliance checkbox for us. We genuinely want to make sure that user data is protected at every level, and that we have thought through every possible risk scenario.")

    heading(doc, "11.1 Types of Data We Process", level=2)
    body(doc, "The platform handles several categories of personal data, and each one requires different levels of protection:")
    data_types = [
        "Account Data — Email addresses, hashed passwords, login timestamps, IP addresses. This is standard authentication data that every web application collects.",
        "Profile Data — Display names, avatar images, language preferences. Relatively low sensitivity but still personal information.",
        "Health Data — This is the most sensitive category. It includes mood scores, mood notes, clinical session notes, and medical records. Any breach of this data could have serious consequences for users.",
        "Communication Data — Encrypted chat messages between patients and clinicians, and AI chatbot conversation transcripts. These conversations are extremely private.",
        "Behavioural Data — Session durations, platform usage patterns, resource view counts. Used for analytics but must be anonymised to prevent individual identification.",
        "Emergency Data — Emergency button activation logs and crisis interaction records. Sensitive because they indicate moments of acute distress.",
        "Consent Data — Data sharing preferences and consent change history. Important for regulatory compliance.",
    ]
    for d in data_types:
        bullet(doc, d)

    heading(doc, "11.2 Risk Assessment Matrix", level=2)
    body(doc, "We identified 10 major privacy risks and assessed each one for severity, defined mitigation measures, and evaluated the residual risk after those measures are in place:")
    table(doc, ["Risk Description", "Initial\nSeverity", "Mitigation Measures", "Residual\nRisk"], [
        ["Unauthorised access to health data", "High", "RBAC enforcement, JWT authentication, consent-based access control (FR17, FR65)", "Low"],
        ["Data breach of personal information", "High", "AES-256 encryption at rest, bcrypt password hashing, HTTPS/TLS 1.2+ in transit (FR23, NFRS03)", "Low"],
        ["Unintended data retention beyond necessity", "Medium", "Right to deletion (FR8d), defined retention policies (FR19), data anonymisation on deletion", "Low"],
        ["Unauthorised data sharing with third parties", "High", "Granular consent management (FR54–55), immediate access revocation on consent withdrawal", "Low"],
        ["Session hijacking or token theft", "Medium", "JWT tokens expire after 1 hour, auto-logout after 15 minutes inactivity (FR62, NFRS02)", "Low"],
        ["SQL injection and XSS attacks", "High", "All inputs validated and sanitised (FR63), parameterised database queries (NFRS05)", "Low"],
        ["Data loss due to system failure", "Medium", "Automated daily backups at midnight (FR45), restoration within 1 hour SLA (FR46)", "Low"],
        ["Excessive API access or DDoS attempts", "Medium", "Rate limiting at 100 requests per IP per minute with HTTP 429 response (NFRS06)", "Low"],
        ["Emergency data misuse or unauthorised access", "Medium", "All emergency activations logged (FR40), complete audit trails for every access (FR22)", "Low"],
        ["Non-compliance with privacy regulations", "High", "GDPR-aligned design, privacy by design principles, granular consent management (NFRC01–04)", "Low"],
    ], col_widths=[2.0, 0.6, 3.0, 0.6])

    heading(doc, "11.3 Privacy Controls Implementation", level=2)
    table(doc, ["Control Area", "How We Have Implemented It"], [
        ["Encryption at Rest", "AES-256 encryption for all sensitive data — chat messages, clinical records, and mood notes"],
        ["Encryption in Transit", "HTTPS with TLS 1.2+ enforced on all API endpoints — no unencrypted connections allowed"],
        ["Password Security", "bcrypt hashing with 10 salt rounds — even if the database is compromised, passwords cannot be reversed"],
        ["Authentication", "JWT tokens expire after 1 hour, OTP verification for new accounts, support for multi-factor authentication"],
        ["Access Control", "Role-Based Access Control with three roles (Patient, Clinician, Admin) — each role has strictly defined permissions"],
        ["Consent Management", "Users can grant or revoke consent for data sharing at any time, with changes taking effect immediately"],
        ["Audit Trails", "Every database transaction, record access, and consent change is logged with timestamp, user ID, and IP address"],
        ["Data Minimisation", "We only collect data that is strictly necessary — no excessive profiling or unnecessary data harvesting"],
        ["Right to Deletion", "Users can request complete account deletion with data anonymisation processed within 24 hours"],
        ["Backup & Recovery", "Daily automated backups at midnight with 7-day retention and a 1-hour restoration SLA"],
    ], col_widths=[1.6, 5.0])
    add_footer_text(doc, "Mental Health Support Platform — Section 11 DPIA")
    save(doc, "24_DPIA.docx")


def gen_25_test_plan():
    doc = new_doc()
    add_section_header(doc, "Section 12", "Test and Implementation Plan", "How We Will Test and Deploy")
    heading(doc, "12. Test and Implementation Plan")
    body(doc, "Testing is not something we are leaving for the last week — it is built into our entire development process. We follow a multi-layered testing strategy that catches bugs at every level, from individual functions to complete user workflows.")

    heading(doc, "12.1 Testing Strategy", level=2)
    body(doc, "Our testing approach has six layers, each serving a different purpose:")
    table(doc, ["Test Type", "Tool", "What It Covers", "Acceptance Criteria"], [
        ["Unit Testing", "Jest", "Individual functions, API endpoint handlers, React component rendering", "Minimum 80% code coverage across the entire codebase"],
        ["Integration Testing", "Jest + Supertest", "API endpoint chains (e.g., register → verify → login), database operations, authentication flows", "All critical integration paths pass without errors"],
        ["End-to-End Testing", "Cypress", "Complete user workflows from the browser — registration, mood tracking, appointment booking, messaging", "All core user journeys complete successfully"],
        ["Security Testing", "OWASP ZAP + Manual", "SQL injection attempts, XSS attacks, authentication bypass, rate limit enforcement", "Zero critical or high severity vulnerabilities"],
        ["Performance Testing", "Artillery / k6", "Page load times, API response times under load, concurrent user handling", "All NFR performance targets met (page < 3s, API < 500ms, 50 users)"],
        ["User Acceptance", "Manual Testing", "Stakeholder validation of all features against the requirements specification", "All 75 functional requirements verified and approved"],
    ], col_widths=[1.1, 1.1, 2.5, 1.8])

    heading(doc, "12.2 Key Test Cases", level=2)
    body(doc, "We have defined 20 critical test cases that cover the most important functionalities. Each test case is directly linked to one or more functional requirements:")
    table(doc, ["TC#", "Test Scenario", "FR", "Procedure and Expected Result", "Status"], [
        ["TC001", "User Registration", "FR1", "Submit valid registration data → Account created, OTP email sent", "Pass"],
        ["TC002", "Duplicate Email", "FR1", "Register with existing email → Error: 'Email already registered'", "Pass"],
        ["TC003", "OTP Verification", "FR2", "Enter correct OTP within 5 seconds → Account marked as verified", "Pass"],
        ["TC004", "Account Lockout", "FR3", "Enter wrong password 5 times → Account locked, error message shown", "Pass"],
        ["TC005", "Password Reset", "FR5", "Request reset, use valid token within 10 minutes → Password updated", "Pass"],
        ["TC006", "Mood Submission", "FR27", "Submit mood score (1–10) with notes → Entry recorded with timestamp", "Pass"],
        ["TC007", "Duplicate Mood Block", "FR28", "Submit second mood within 1 hour → Rejected with message", "Pass"],
        ["TC008", "Risk Alert", "FR30", "Log mood < 2 for 3 consecutive days → Risk alert notification created", "Pass"],
        ["TC009", "Book Appointment", "FR9", "Select available slot → Appointment created, confirmation sent", "Pass"],
        ["TC010", "Slot Conflict", "FR10", "Book an already-taken slot → Rejected: 'Slot unavailable'", "Pass"],
        ["TC011", "Emergency Button", "FR38", "Press emergency button → Crisis info displayed within 1 second", "Pass"],
        ["TC012", "Message Encryption", "FR23", "Send chat message → Stored in database as AES-256 encrypted text", "Pass"],
        ["TC013", "Chatbot Response", "FR32", "Send query to chatbot → AI response received within 3 seconds", "Pass"],
        ["TC014", "Chatbot Escalation", "FR33", "Trigger low-confidence response → Conversation escalated to clinician", "Pass"],
        ["TC015", "Consent Revoke", "FR54", "Revoke data sharing consent → All data access immediately blocked", "Pass"],
        ["TC016", "Auto-Logout", "FR62", "Leave session idle for 15 minutes → Automatically logged out", "Pass"],
        ["TC017", "RBAC Check", "FR65", "Patient tries to access admin route → HTTP 403 Forbidden returned", "Pass"],
        ["TC018", "SQL Injection", "FR63", "Inject SQL string in input field → Input sanitised, no DB impact", "Pass"],
        ["TC019", "File Size Limit", "FR73", "Upload file larger than 5MB → Rejected with size limit message", "Pass"],
        ["TC020", "Dashboard Speed", "FR50", "Load main dashboard → Full render completed within 3 seconds", "Pass"],
    ], col_widths=[0.5, 1.2, 0.4, 3.3, 0.5])

    heading(doc, "12.3 Implementation Plan", level=2)
    body(doc, "Our implementation follows a phased approach, with each phase building on top of the previous one:")
    table(doc, ["Phase", "Focus Area", "Detailed Activities"], [
        ["Phase 1\nWeek 7", "Backend Foundation", "Setting up Express.js project structure, PostgreSQL schema migration using Prisma, implementing JWT + OTP authentication module, building user CRUD APIs, and basic middleware (error handling, logging, CORS)"],
        ["Phase 2\nWeek 8", "Core Feature APIs", "Mood tracking endpoints, appointment scheduling with conflict detection, messaging with AES-256 encryption, chatbot integration with OpenAI API, emergency support endpoints, and notification engine setup with Twilio and SendGrid"],
        ["Phase 3\nWeek 9", "Frontend Development", "React component library creation, dashboard page with Chart.js, mood tracker page, messaging UI with Socket.io client, appointment booking interface, admin panel, and responsive layout with Tailwind CSS"],
        ["Phase 4\nWeek 10", "Integration", "Connecting frontend to backend APIs, Socket.io real-time chat activation, Google Calendar synchronisation, notification engine end-to-end testing, and comprehensive integration testing with Supertest"],
        ["Phase 5\nWeek 11", "Testing & Deploy", "Running full Jest unit test suite, Cypress E2E tests, security audit with OWASP guidelines, performance testing with Artillery, cloud deployment to Render/AWS, and CI/CD pipeline with GitHub Actions"],
        ["Phase 6\nWeek 12", "Demo & Handover", "Final demonstration to evaluators, documentation completion, source code handover, deployment documentation, and lessons learned report"],
    ], col_widths=[0.8, 1.2, 4.6])
    add_footer_text(doc, "Mental Health Support Platform — Section 12 Test and Implementation Plan")
    save(doc, "25_Test_Implementation_Plan.docx")


def gen_26_feedback():
    doc = new_doc()
    add_section_header(doc, "Section 13", "Responses on Feedback\nfrom Supervisor", "How We Acted on Review Comments")
    heading(doc, "13. Responses on Feedback from Supervisor")
    body(doc, "Throughout this project, we have been receiving regular feedback from our academic supervisor, Mr. Syed Altaf, during weekly review sessions. We have also received input from our industry sponsor, Mr. Nabin Singh from Skillup Labs. This section documents the key feedback points and how we addressed each one of them.")
    body(doc, "We believe that feedback is not criticism — it is an opportunity to make the product better. Every piece of feedback we received actually improved the quality of our deliverables significantly.")

    table(doc, ["Week", "Feedback Received", "Our Response and Action Taken"], [
        ["Week 2", "The scope boundaries are not clear enough. It is not obvious what is included in the prototype and what is being left out.",
         "We completely rewrote Section 1.3 (Scope and Limitations) to clearly distinguish between what the prototype will include and what is outside scope. We explicitly stated that video consultation, native mobile app, and real clinical diagnosis are all out of scope for this iteration."],
        ["Week 3", "The functional requirements are too vague. They need to be measurable and testable — otherwise how will you know if you have met them?",
         "This was really valuable feedback. We went back and revised all 75 functional requirements to include specific measurable criteria. For example, instead of saying 'OTP should be fast', we specified 'OTP delivery within 5 seconds'. Each FR now has a corresponding test case with clear pass/fail criteria."],
        ["Week 4", "The architecture diagram does not clearly show the separation of concerns. It is hard to tell which components belong to which tier.",
         "We redesigned the entire system architecture diagram with clear colour-coded sections for each tier (Presentation, Application, Data). Added specific technology labels to every component and included data flow arrows showing communication protocols (HTTPS, WSS, SQL)."],
        ["Week 5", "The Data Flow Diagrams are mixing notations. Please stick to one method consistently — Gane & Sarson is preferred.",
         "We redid all DFDs from scratch using strictly Gane & Sarson notation — rounded rectangles for processes, open-ended rectangles for data stores, and rectangles for external entities. Created a consistent visual language across all 17 DFD diagrams (1 Context, 1 Level 1, 15 Level 2)."],
        ["Week 5", "The ERD needs to be database-ready. It should be normalised properly and ready for direct implementation.",
         "We redesigned the ERD with proper normalisation, UUIDs as primary keys, explicit foreign key relationships, and added a comprehensive data dictionary with column-level detail including data types, constraints, and descriptions for every single field."],
        ["Week 6", "Security considerations should not just be in one section — they need to be woven throughout the entire design.",
         "We integrated security throughout — added the DPIA (Section 11) with a full risk assessment, included security-specific NFRs (NFRS01-06), ensured all security-related FRs (FR23, FR54-55, FR62-66) appear in the DFDs with proper data store references, and added security test cases (TC012, TC017, TC018) to the test plan."],
    ], col_widths=[0.6, 2.2, 3.8])

    body(doc, "We are grateful to Mr. Syed Altaf for his thorough and constructive feedback. It has genuinely made this project much stronger than it would have been otherwise.")
    add_footer_text(doc, "Mental Health Support Platform — Section 13 Responses on Feedback")
    save(doc, "26_Supervisor_Feedback.docx")


def gen_27_references():
    doc = new_doc()
    add_section_header(doc, "References", "References", "Sources and Citations Used in This Report")
    heading(doc, "References")
    body(doc, "The following sources were referred to during the research, design, and documentation phases of this project. All references follow the APA referencing style as required by Kent Institute of Technology Australia.")

    refs = [
        "[1]  Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education. — Used as the primary reference for software engineering principles, SDLC methodology, and requirements engineering best practices.",
        "[2]  Pressman, R. S. & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill. — Referenced for system design patterns, testing strategies, and project management techniques.",
        "[3]  IEEE (1998). IEEE Std 830-1998: Recommended Practice for Software Requirements Specifications. IEEE. — Followed the IEEE standard format for structuring our Software Requirements Specification document.",
        "[4]  Gane, C. & Sarson, T. (1979). Structured Systems Analysis: Tools and Techniques. Prentice-Hall. — The definitive reference for the Gane & Sarson DFD notation used throughout our data flow diagrams.",
        "[5]  OWASP Foundation (2021). OWASP Top 10 — 2021. https://owasp.org/Top10/ — Used as the security baseline for identifying and mitigating common web application vulnerabilities.",
        "[6]  React.js Documentation (2024). React — A JavaScript Library for Building User Interfaces. https://react.dev/ — Official documentation for our frontend framework.",
        "[7]  Express.js Documentation (2024). Express — Fast, Unopinionated, Minimalist Web Framework. https://expressjs.com/ — Official documentation for our backend API framework.",
        "[8]  PostgreSQL Documentation (2024). PostgreSQL: The World's Most Advanced Open Source Database. https://www.postgresql.org/docs/ — Database documentation for schema design and query optimisation.",
        "[9]  OpenAI API Documentation (2024). API Reference. https://platform.openai.com/docs/ — Documentation for integrating the AI chatbot with the OpenAI GPT models.",
        "[10] Google Calendar API Documentation (2024). Calendar API Overview. https://developers.google.com/calendar/ — Used for implementing appointment synchronisation with Google Calendar.",
        "[11] Twilio Documentation (2024). Programmable Messaging. https://www.twilio.com/docs/ — Documentation for SMS notification and OTP delivery integration.",
        "[12] SendGrid Documentation (2024). Email API. https://docs.sendgrid.com/ — Documentation for transactional email notification integration.",
        "[13] Socket.IO Documentation (2024). Socket.IO. https://socket.io/docs/ — Documentation for implementing real-time WebSocket communication for chat and notifications.",
        "[14] W3C (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/ — Accessibility standards that our UI design complies with.",
        "[15] Information Commissioner's Office (ICO) (2022). Data Protection Impact Assessments. https://ico.org.uk/ — Guidelines followed for our DPIA methodology and privacy risk assessment.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        r.font.size = Pt(10)
        r.font.name = 'Calibri'
        r.font.color.rgb = GREY_DARK
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    add_footer_text(doc, "Mental Health Support Platform — References")
    save(doc, "27_References.docx")


def gen_28_appendix():
    doc = new_doc()
    add_section_header(doc, "Appendix", "Appendix", "Table of Figures and Table of Tables")
    heading(doc, "Appendix")

    heading(doc, "A. Table of Figures", level=2)
    body(doc, "A complete listing of all diagrams and figures included across the 29 sections of this report:")
    table(doc, ["Figure #", "Description", "Section"], [
        ["Figure 1.1", "Organisation Chart — Mel23 Tech Solution", "1.6"],
        ["Figure 1.2", "Work Breakdown Structure", "1.10"],
        ["Figure 3.1", "Three-Tier System Architecture Diagram", "3.1"],
        ["Figure 4.1", "Login / Registration Page Wireframe", "4.1"],
        ["Figure 4.2", "User Dashboard Wireframe", "4.1"],
        ["Figure 4.3", "Mood Tracker Page Wireframe", "4.1"],
        ["Figure 4.4", "Secure Messaging Interface Wireframe", "4.1"],
        ["Figure 4.5", "Appointments Page Wireframe", "4.1"],
        ["Figure 4.6", "Admin Dashboard Panel Wireframe", "4.1"],
        ["Figure 5.1", "Patient / User Use Case Diagram", "5"],
        ["Figure 5.2", "Clinician / Doctor Use Case Diagram", "5"],
        ["Figure 5.3", "Administrator Use Case Diagram", "5"],
        ["Figure 6.1", "Level 0 Context Diagram (Gane & Sarson)", "6"],
        ["Figure 7.1", "Level 1 System Data Flow Diagram", "7.1"],
        ["Figure 7.2.1", "Level 2.1 — Authentication & Account Management DFD", "7.2"],
        ["Figure 7.2.2", "Level 2.2 — Profile Management DFD", "7.2"],
        ["Figure 7.2.3", "Level 2.3 — Appointment Management DFD", "7.2"],
        ["Figure 7.2.4", "Level 2.4 — Clinical Records Management DFD", "7.2"],
        ["Figure 7.2.5", "Level 2.5 — Messaging System DFD", "7.2"],
        ["Figure 7.2.6", "Level 2.6 — Mood Tracking & Analysis DFD", "7.2"],
        ["Figure 7.2.7", "Level 2.7 — AI Chatbot System DFD", "7.2"],
        ["Figure 7.2.8", "Level 2.8 — Emergency Support DFD", "7.2"],
        ["Figure 7.2.9", "Level 2.9 — Notification Engine DFD", "7.2"],
        ["Figure 7.2.10", "Level 2.10 — Content & Resource Management DFD", "7.2"],
        ["Figure 7.2.11", "Level 2.11 — Analytics & Reporting DFD", "7.2"],
        ["Figure 7.2.12", "Level 2.12 — Administration & Access Control DFD", "7.2"],
        ["Figure 7.2.13", "Level 2.13 — System Maintenance & Backup DFD", "7.2"],
        ["Figure 7.2.14", "Level 2.14 — Privacy & Compliance DFD", "7.2"],
        ["Figure 7.2.15", "Level 2.15 — UI & Localisation DFD", "7.2"],
        ["Figure 8.1", "SD1 — User Registration & Login Sequence Diagram", "8"],
        ["Figure 8.2", "SD2 — Appointment Booking Flow Sequence Diagram", "8"],
        ["Figure 8.3", "SD3 — Mood Entry & Risk Alert Sequence Diagram", "8"],
        ["Figure 8.4", "SD4 — AI Chatbot Interaction Sequence Diagram", "8"],
        ["Figure 8.5", "SD5 — Emergency Button Activation Sequence Diagram", "8"],
        ["Figure 8.6", "SD6 — Secure Messaging Sequence Diagram", "8"],
        ["Figure 9.1", "Entity Relationship Diagram", "9"],
        ["Figure 10.1", "UML Class Diagram", "10"],
    ], col_widths=[1.0, 4.0, 0.6])

    heading(doc, "B. Table of Tables", level=2)
    table(doc, ["Table #", "Description", "Section"], [
        ["Table 1.1", "Company Information", "1.5"],
        ["Table 1.2", "Roles and Responsibilities", "1.7"],
        ["Table 1.3", "Key Deliverables", "1.8"],
        ["Table 1.4", "Project Budget Estimates", "1.9"],
        ["Table 1.5", "WBS Phase Summary", "1.10"],
        ["Table 1.6", "Project Timeline / Gantt Chart", "1.11"],
        ["Table 2.1", "Functional Requirements (FR1–FR75)", "2.1"],
        ["Table 2.2", "Non-Functional Requirements (6 categories)", "2.2"],
        ["Table 3.1", "Technology Stack", "3.2"],
        ["Table 7.1", "Level 1 Process Descriptions", "7.1"],
        ["Table 7.2", "Data Store Definitions", "7.1"],
        ["Table 9.1–9.11", "Data Dictionary (11 entity tables)", "9"],
        ["Table 11.1", "DPIA Risk Assessment Matrix", "11"],
        ["Table 11.2", "Privacy Controls Implementation", "11"],
        ["Table 12.1", "Testing Strategy", "12"],
        ["Table 12.2", "Test Cases Summary (20 cases)", "12"],
        ["Table 12.3", "Implementation Plan (6 phases)", "12"],
        ["Table 13.1", "Supervisor Feedback and Responses", "13"],
    ], col_widths=[1.0, 4.0, 0.6])
    add_footer_text(doc, "Mental Health Support Platform — Appendix")
    save(doc, "28_Appendix.docx")


# Not in the 29 count as per the user's request, but we need exactly 29.
# Recounting: 01-11 (11 from Section 1) + 12-13 (2 from Section 2) + 14-15 (2 from Section 3)
# + 16 (1 from Section 4) + 17-26 (Sections 5-13 = items but 17=UC, 18=Context, 19=DFD1, 20=DFD2,
#   21=Seq, 22=ERD, 23=Class, 24=DPIA, 25=Test, 26=Feedback = 10 + but Section 7 has 2 sub = 11)
# + 27-28 (References + Appendix) = 2
# Total = 11 + 2 + 2 + 1 + 11 + 2 = 29 ✓


# ════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  Generating 29 Report Section Files")
    print("  Output: Report_Sections/")
    print("=" * 60)

    generators = [
        gen_00_cover_toc,
        gen_01_purpose, gen_02_objective, gen_03_scope, gen_04_assumptions,
        gen_05_company_profile, gen_06_org_chart, gen_07_roles, gen_08_deliverables,
        gen_09_budget, gen_10_wbs, gen_11_gantt,
        gen_12_functional_req, gen_13_nonfunctional_req,
        gen_14_architecture, gen_15_tech_stack,
        gen_16_wireframes,
        gen_17_use_case, gen_18_context_diagram, gen_19_dfd_level1, gen_20_dfd_level2,
        gen_21_sequence, gen_22_erd, gen_23_class_diagram,
        gen_24_dpia, gen_25_test_plan, gen_26_feedback,
        gen_27_references, gen_28_appendix,
    ]

    print(f"\n  Total generators: {len(generators)}")
    print()

    for gen_func in generators:
        gen_func()

    print(f"\n{'=' * 60}")
    print(f"  ✅ All {len(generators)} files generated in: {OUT_DIR}")
    print(f"{'=' * 60}")
